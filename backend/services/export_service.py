"""
Ephemeral export service for LACUNEX AI.
Generates .pdf, .docx, and .xlsx files in-memory.
No data is stored or logged.
"""

import io
import re
from datetime import datetime

# ── LaTeX stripping patterns ──────────────────────────────────────────────
_LATEX_FRAC = re.compile(r'\\frac\{([^}]*)\}\{([^}]*)\}')
_LATEX_SQRT = re.compile(r'\\sqrt\{([^}]*)\}')
_LATEX_INLINE = re.compile(r'\$([^$]+)\$')
_LATEX_DISPLAY = re.compile(r'\$\$(.+?)\$\$', re.DOTALL)
_LATEX_COMMANDS = {
    r'\\alpha': 'alpha', r'\\beta': 'beta', r'\\gamma': 'gamma',
    r'\\delta': 'delta', r'\\epsilon': 'epsilon', r'\\theta': 'theta',
    r'\\lambda': 'lambda', r'\\mu': 'mu', r'\\pi': 'pi',
    r'\\sigma': 'sigma', r'\\phi': 'phi', r'\\psi': 'psi',
    r'\\omega': 'omega', r'\\infty': 'infinity',
    r'\\sum': 'SUM', r'\\prod': 'PRODUCT',
    r'\\int': 'INTEGRAL', r'\\partial': 'd',
    r'\\nabla': 'nabla', r'\\cdot': '*', r'\\times': 'x',
    r'\\rightarrow': '->', r'\\leftarrow': '<-',
    r'\\Rightarrow': '=>', r'\\Leftarrow': '<=',
    r'\\leq': '<=', r'\\geq': '>=', r'\\neq': '!=',
    r'\\approx': '~=', r'\\equiv': '===',
    r'\\in': 'in', r'\\notin': 'not in',
    r'\\subset': 'subset', r'\\cup': 'U', r'\\cap': 'n',
    r'\\forall': 'for all', r'\\exists': 'exists',
    r'\\quad': ' ', r'\\qquad': '  ',
    r'\\text': '', r'\\mathrm': '', r'\\mathbf': '',
}


def _strip_latex(text: str) -> str:
    """Strip LaTeX notation and convert to readable plain text."""
    if not text or '$' not in text and '\\' not in text:
        return text
    # Display math first
    text = _LATEX_DISPLAY.sub(r'\1', text)
    # Fractions: \frac{a}{b} -> a/b
    text = _LATEX_FRAC.sub(r'\1/\2', text)
    # Square roots: \sqrt{x} -> sqrt(x)
    text = _LATEX_SQRT.sub(r'sqrt(\1)', text)
    # Ket/bra notation: |\psi\rangle -> |psi>
    text = re.sub(r'\|\\([a-zA-Z]+)\\rangle', r'|\1>', text)
    text = re.sub(r'\\langle\\([a-zA-Z]+)\|', r'<\1|', text)
    text = re.sub(r'\\rangle', '>', text)
    text = re.sub(r'\\langle', '<', text)
    # Named commands
    for pattern, replacement in _LATEX_COMMANDS.items():
        text = re.sub(pattern, replacement, text)
    # Remaining backslash commands: \command{arg} -> arg
    text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
    # Remaining backslash commands without args: \command -> ""
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    # Inline math: $...$ -> content
    text = _LATEX_INLINE.sub(r'\1', text)
    # Clean up braces left over
    text = text.replace('{', '').replace('}', '')
    # Clean up multiple spaces
    text = re.sub(r'  +', ' ', text)
    return text.strip()

BRAND_FOOTER = "\u00a9 Generated from Lacunex AI by Shasradha Karmakar"


def _ts() -> str:
    return datetime.utcnow().strftime("%B %d, %Y at %H:%M UTC")


def _clean(text: str) -> str:
    """Normalize line endings and remove null bytes."""
    return text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n") if text else ""


def _pdf_safe(text: str) -> str:
    """
    Convert text to ASCII-safe string for FPDF Helvetica (Latin-1 font).
    Strips LaTeX first, then maps Unicode to ASCII equivalents.
    """
    # Strip LaTeX before Unicode mapping
    text = _strip_latex(text)
    replacements = {
        "\u2018": "'", "\u2019": "'",          # smart single quotes
        "\u201c": '"', "\u201d": '"',          # smart double quotes
        "\u2013": "-", "\u2014": "--",         # en-dash, em-dash
        "\u2022": "*", "\u00b7": "*",          # bullet
        "\u2026": "...",                       # ellipsis
        "\u00a9": "(c)", "\u00ae": "(R)",      # copyright, registered
        "\u00b0": "deg", "\u00b1": "+/-",      # degree, plus-minus
        "\u00d7": "x", "\u00f7": "/",          # multiply, divide
        "\u2212": "-", "\u00ac": "not",        # minus sign, not
        "\u2260": "!=", "\u2264": "<=",        # not-equal, leq
        "\u2265": ">=", "\u221e": "inf",       # geq, infinity
        "\u03b1": "alpha", "\u03b2": "beta",   # Greek letters
        "\u03b3": "gamma", "\u03b4": "delta",
        "\u03b5": "epsilon", "\u03b8": "theta",
        "\u03c0": "pi", "\u03bb": "lambda",
        "\u03c3": "sigma", "\u03c6": "phi",
        "\u03c8": "psi", "\u03c9": "omega",
        "\u2192": "->", "\u2190": "<-",        # arrows
        "\u21d2": "=>", "\u21d0": "<=",        # double arrows
        "\u2713": "OK", "\u2717": "X",         # check, cross
        "\u2248": "~=", "\u2261": "===",       # approx, equiv
    }
    out = []
    for ch in text:
        if ch in replacements:
            out.append(replacements[ch])
            continue
        try:
            ch.encode("latin-1")
            out.append(ch)
        except (UnicodeEncodeError, UnicodeDecodeError):
            out.append("?")
    return "".join(out)


_MD_BOLD = re.compile(r"\*{2,3}(.+?)\*{2,3}", re.DOTALL)
_MD_ITALIC = re.compile(r"\*(.+?)\*", re.DOTALL)
_MD_CODE_INLINE = re.compile(r"`([^`]+)`")
_MD_LINK = re.compile(r"\[([^\]]+)\]\([^)]+\)")
_MD_HEADING = re.compile(r"^#{1,6}\s+", re.MULTILINE)
_MD_HR = re.compile(r"^[-*_]{3,}$")


def _strip_md(text: str) -> str:
    """
    Fully strip markdown and LaTeX to produce clean plain text.
    Processes line-by-line for accuracy.
    """
    # Strip LaTeX first
    text = _strip_latex(text)
    lines = []
    for line in text.split("\n"):
        # Headings
        line = re.sub(r"^#{1,6}\s+", "", line)
        # Bold/italic markers
        line = _MD_BOLD.sub(r"\1", line)
        line = _MD_ITALIC.sub(r"\1", line)
        # Inline code
        line = _MD_CODE_INLINE.sub(r"\1", line)
        # Links
        line = _MD_LINK.sub(r"\1", line)
        # HR → blank
        if _MD_HR.match(line.strip()):
            line = ""
        # Numbered / unordered lists → keep text cleanly
        line = re.sub(r"^(\s*)[-*]\s+", r"\1", line)
        line = re.sub(r"^(\s*)\d+\.\s+", r"\1", line)
        lines.append(line)
    return "\n".join(lines)


def _parse_blocks(content: str):
    """
    Parse content into a list of block dicts:
      { type: heading|bullet|text|code|blank|hr, level: int, text: str, lines: [...] }
    """
    blocks = []
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        ln = lines[i]
        stripped = ln.strip()

        # Code fence
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append({"type": "code", "lines": code_lines})
            i += 1
            continue

        # Blank
        if not stripped:
            blocks.append({"type": "blank"})
            i += 1
            continue

        # HR
        if _MD_HR.match(stripped):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue

        # Bullet (unordered)
        m = re.match(r"^[-*]\s+(.*)", stripped)
        if m:
            blocks.append({"type": "bullet", "text": m.group(1).strip()})
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            blocks.append({"type": "numbered", "number": int(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue

        # Regular text
        blocks.append({"type": "text", "text": stripped})
        i += 1

    return blocks


def _inline_runs(text: str):
    """Split text into runs with bold/italic/code markers parsed."""
    runs = []
    pattern = re.compile(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`[^`]+`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            runs.append({"text": text[last:m.start()], "bold": False, "italic": False, "code": False})
        token = m.group(0)
        if token.startswith("***"):
            runs.append({"text": token[3:-3], "bold": True, "italic": True, "code": False})
        elif token.startswith("**"):
            runs.append({"text": token[2:-2], "bold": True, "italic": False, "code": False})
        elif token.startswith("*"):
            runs.append({"text": token[1:-1], "bold": False, "italic": True, "code": False})
        else:
            runs.append({"text": token[1:-1], "bold": False, "italic": False, "code": True})
        last = m.end()
    if last < len(text):
        runs.append({"text": text[last:], "bold": False, "italic": False, "code": False})
    return runs or [{"text": text, "bold": False, "italic": False, "code": False}]


# ─────────────────────────────────────────────────────────────────────────────
#  PDF  — fpdf2, proper Unicode handling
# ─────────────────────────────────────────────────────────────────────────────

def generate_pdf(title: str, messages: list[dict], model_name: str | None = None) -> bytes:
    from fpdf import FPDF

    class LacunexPDF(FPDF):
        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(120, 120, 130)
            self.cell(0, 5, _pdf_safe(BRAND_FOOTER), align="C")

    pdf = LacunexPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    L = 20
    R = 20
    PW = 210 - L - R

    pdf.set_left_margin(L)
    pdf.set_right_margin(R)
    pdf.set_top_margin(20)

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(20, 20, 40)
    pdf.multi_cell(PW, 8, _pdf_safe(_clean(title)), align="L")
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(140, 140, 150)
    pdf.cell(PW, 5, _pdf_safe(f"Exported on {_ts()}"), ln=True)

    pdf.set_draw_color(200, 200, 220)
    pdf.ln(4)
    pdf.line(L, pdf.get_y(), 210 - R, pdf.get_y())
    pdf.ln(6)

    for msg in messages:
        role = msg.get("role", "")
        raw = _clean(msg.get("content", ""))
        if not raw:
            continue

        is_user = role == "user"
        ai_label = f"Lacunex AI ({model_name})" if model_name else "Lacunex AI"

        pdf.set_font("Helvetica", "B", 9)
        if is_user:
            pdf.set_text_color(40, 110, 210)
        else:
            pdf.set_text_color(110, 60, 210)
            
        pdf.cell(PW, 6, "You" if is_user else _pdf_safe(ai_label), ln=True)
        pdf.ln(1)

        blocks = _parse_blocks(raw)
        for block in blocks:
            btype = block.get("type")

            if btype == "blank":
                pdf.ln(2)

            elif btype == "hr":
                pdf.set_draw_color(220, 220, 230)
                pdf.line(L, pdf.get_y(), 210 - R, pdf.get_y())
                pdf.ln(3)

            elif btype == "heading":
                lvl = block["level"]
                sizes = {1: 14, 2: 12, 3: 11, 4: 10}
                pdf.set_font("Helvetica", "B", sizes.get(lvl, 10))
                pdf.set_text_color(30, 30, 50)
                pdf.ln(2)
                pdf.multi_cell(PW, 6, _pdf_safe(block["text"]), align="L")
                pdf.ln(1)

            elif btype == "bullet":
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 60)
                clean_text = _strip_md(block["text"])
                
                # Render bullet point at fixed X
                pdf.set_x(L + 2)
                pdf.cell(5, 5.5, _pdf_safe("\u2022"))
                
                # Render text with consistent indentation
                pdf.set_left_margin(L + 7)
                pdf.set_x(L + 7)
                pdf.multi_cell(PW - 7, 5.5, _pdf_safe(clean_text), align="L")
                pdf.set_left_margin(L)

            elif btype == "numbered":
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 60)
                clean_text = _strip_md(block["text"])
                num_str = f"{block['number']}."
                
                # Render number at fixed X
                pdf.set_x(L + 2)
                pdf.cell(7, 5.5, _pdf_safe(num_str))
                
                # Render text with consistent indentation
                pdf.set_left_margin(L + 9)
                pdf.set_x(L + 9)
                pdf.multi_cell(PW - 9, 5.5, _pdf_safe(clean_text), align="L")
                pdf.set_left_margin(L)

            elif btype == "code":
                pdf.set_fill_color(248, 248, 252)
                pdf.set_draw_color(220, 220, 240)
                pdf.set_font("Courier", "", 9)
                pdf.set_text_color(60, 50, 140)
                pdf.ln(2)
                for code_line in block.get("lines", []):
                    safe = _pdf_safe(code_line if code_line else " ")
                    # small indent for code block
                    pdf.set_x(L + 2)
                    pdf.multi_cell(PW - 4, 5, " " + safe, fill=True, border=0, align="L")
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 60)
                pdf.ln(3)

            elif btype == "text":
                clean_text = _strip_md(block["text"])
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 60)
                pdf.multi_cell(PW, 5.5, _pdf_safe(clean_text), align="L")

        pdf.ln(6)
        pdf.set_draw_color(230, 230, 240)
        pdf.line(L, pdf.get_y(), 210 - R, pdf.get_y())
        pdf.ln(6)

    return bytes(pdf.output())


# ─────────────────────────────────────────────────────────────────────────────
#  DOCX  — proper Word styles, no raw markdown
# ─────────────────────────────────────────────────────────────────────────────

def generate_docx(title: str, messages: list[dict], model_name: str | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ── Helpers ────────────────────────────────────────────────────────────
    def add_hr():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCDD")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

    def shade_para(p, color_hex: str):
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color_hex)
        pPr.append(shd)

    def add_runs(para, text: str, base_size: float = 10.5, base_color=None):
        """Add inline-parsed runs to a paragraph."""
        if base_color is None:
            base_color = RGBColor(22, 22, 45)
        for r in _inline_runs(text):
            run = para.add_run(r["text"])
            run.bold = r["bold"]
            run.italic = r["italic"]
            run.font.size = Pt(base_size)
            if r["code"]:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(80, 55, 180)
            else:
                run.font.color.rgb = base_color

    # ── Title ────────────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tr = tp.add_run(_clean(title))
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(12, 12, 38)
    tp.paragraph_format.space_after = Pt(4)

    dp = doc.add_paragraph()
    dr = dp.add_run(f"Exported on {_ts()}")
    dr.font.size = Pt(8.5)
    dr.font.color.rgb = RGBColor(145, 145, 168)
    dr.italic = True
    dp.paragraph_format.space_after = Pt(12)

    add_hr()

    # ── Messages ─────────────────────────────────────────────────────────────
    for msg in messages:
        role = msg.get("role", "")
        raw = _clean(msg.get("content", ""))
        if not raw:
            continue

        is_user = role == "user"

        # Sender label
        ai_label = f"Lacunex AI ({model_name})" if model_name else "Lacunex AI"
        lp = doc.add_paragraph()
        lr = lp.add_run("You" if is_user else ai_label)
        lr.bold = True
        lr.font.size = Pt(8.5)
        lr.font.color.rgb = RGBColor(40, 120, 215) if is_user else RGBColor(100, 65, 215)
        lp.paragraph_format.space_before = Pt(10)
        lp.paragraph_format.space_after = Pt(4)

        # Render blocks
        blocks = _parse_blocks(raw)
        for block in blocks:
            btype = block.get("type")

            if btype == "blank":
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(0)
                sp.paragraph_format.space_after = Pt(1)

            elif btype == "hr":
                add_hr()

            elif btype == "heading":
                lvl = block["level"]
                sizes = {1: 15, 2: 13, 3: 12, 4: 11}
                hp = doc.add_paragraph()
                hr_ = hp.add_run(block["text"])
                hr_.bold = True
                hr_.font.size = Pt(sizes.get(lvl, 11))
                hr_.font.color.rgb = RGBColor(14, 14, 40)
                hp.paragraph_format.space_before = Pt(8)
                hp.paragraph_format.space_after = Pt(3)

            elif btype == "bullet":
                bp = doc.add_paragraph(style="List Bullet")
                add_runs(bp, block["text"])
                bp.paragraph_format.left_indent = Inches(0.25)
                bp.paragraph_format.first_line_indent = Inches(-0.25)
                bp.paragraph_format.space_after = Pt(2)

            elif btype == "numbered":
                np_ = doc.add_paragraph(style="List Number")
                add_runs(np_, block["text"])
                np_.paragraph_format.left_indent = Inches(0.25)
                np_.paragraph_format.first_line_indent = Inches(-0.25)
                np_.paragraph_format.space_after = Pt(2)

            elif btype == "code":
                for cl in block.get("lines", []):
                    cp = doc.add_paragraph()
                    cr = cp.add_run(cl if cl else " ")
                    cr.font.name = "Courier New"
                    cr.font.size = Pt(9)
                    cr.font.color.rgb = RGBColor(65, 50, 160)
                    cp.paragraph_format.left_indent = Cm(0.5)
                    cp.paragraph_format.right_indent = Cm(0.5)
                    cp.paragraph_format.space_after = Pt(0)
                    shade_para(cp, "F0F0FA")

                gap = doc.add_paragraph()
                gap.paragraph_format.space_after = Pt(4)

            elif btype == "text":
                p = doc.add_paragraph()
                add_runs(p, block["text"])
                p.paragraph_format.space_after = Pt(3)

        add_hr()

    # ── Footer ────────────────────────────────────────────────────────────────
    fp = doc.add_paragraph()
    fr = fp.add_run(BRAND_FOOTER)
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = RGBColor(155, 155, 175)
    fr.italic = True
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(16)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
#  XLSX  — markdown-free, structured
# ─────────────────────────────────────────────────────────────────────────────

def generate_xlsx(title: str, messages: list[dict], model_name: str | None = None) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "Conversation"

    # Title row
    ws.merge_cells("A1:D1")
    ws["A1"] = _clean(title)
    ws["A1"].font = Font(bold=True, size=14, color="0A0A22")
    ws["A1"].alignment = Alignment(horizontal="left", vertical="center")

    # Subtitle row
    ws.merge_cells("A2:D2")
    ws["A2"] = f"Exported on {_ts()}"
    ws["A2"].font = Font(size=9, color="8282A0", italic=True)
    ws["A2"].alignment = Alignment(horizontal="left")

    ws.append([])  # spacer

    # Headers
    ws.append(["#", "Sender", "Message", "Timestamp"])
    hdr_row = ws.max_row
    hdr_fill = PatternFill("solid", fgColor="4A3FBF")
    for col in range(1, 5):
        cell = ws.cell(row=hdr_row, column=col)
        cell.font = Font(bold=True, color="FFFFFF", size=10)
        cell.fill = hdr_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # Data
    thin = Side(style="thin", color="E0E0EE")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    user_fill = PatternFill("solid", fgColor="EEF3FF")
    ai_fill = PatternFill("solid", fgColor="F2EEFF")
    now = datetime.utcnow().isoformat()

    for idx, msg in enumerate(messages, start=1):
        role = msg.get("role", "user")
        # IMPORTANT: Strip all markdown to clean plain text for Excel
        raw_content = _clean(msg.get("content", ""))
        content = _strip_md(raw_content)
        if not content.strip():
            continue

        ai_label = f"Lacunex AI ({model_name})" if model_name else "Lacunex AI"
        sender = "You" if role == "user" else ai_label
        fill = user_fill if role == "user" else ai_fill
        ts = msg.get("created_at") or now

        ws.append([idx, sender, content, ts])
        data_row = ws.max_row
        for col in range(1, 5):
            cell = ws.cell(row=data_row, column=col)
            cell.fill = fill
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.font = Font(
                bold=(col == 2),
                size=9 if col == 2 else (8 if col == 4 else 10),
                color="828296" if col == 4 else "0A0A22",
            )

    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 13
    ws.column_dimensions["C"].width = 85
    ws.column_dimensions["D"].width = 24
    ws.freeze_panes = "A5"

    ws.append([])
    ws.append(["", "", BRAND_FOOTER])
    footer_cell = ws.cell(row=ws.max_row, column=3)
    footer_cell.font = Font(size=8, color="9090B0", italic=True)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
#  DOCUMENT-MODE EXPORTS — from structured JSON (document_parser output)
# ─────────────────────────────────────────────────────────────────────────────

def _doc_flatten_content(section: dict, depth: int = 0) -> list[dict]:
    """Flatten a section tree into a list of renderable blocks for export."""
    blocks = []
    level = section.get("level", 2)
    heading = section.get("heading", "")

    if heading:
        blocks.append({"type": "heading", "level": level, "text": heading})

    if section.get("content"):
        for line in section["content"].split("\n"):
            stripped = line.strip()
            if not stripped:
                blocks.append({"type": "blank"})
            elif re.match(r"^[-*+]\s+", stripped):
                blocks.append({"type": "bullet", "text": re.sub(r"^[-*+]\s+", "", stripped)})
            elif re.match(r"^\d+\.\s+", stripped):
                m = re.match(r"^(\d+)\.\s+(.*)", stripped)
                if m:
                    blocks.append({"type": "numbered", "number": int(m.group(1)), "text": m.group(2)})
            else:
                blocks.append({"type": "text", "text": stripped})

    for table in section.get("tables", []):
        blocks.append({"type": "table", "data": table})

    for code in section.get("code_blocks", []):
        blocks.append({"type": "code", "lines": code.get("code", "").split("\n"), "language": code.get("language", "")})

    for highlight in section.get("highlights", []):
        h_type = highlight.get("type", "key_point")
        blocks.append({"type": "callout", "callout_type": h_type, "text": highlight.get("text", "")})

    for diagram in section.get("diagrams", []):
        blocks.append({"type": "diagram", "title": diagram.get("title", "Diagram"), "code": diagram.get("code", "")})

    for sub in section.get("subsections", []):
        blocks.extend(_doc_flatten_content(sub, depth + 1))

    return blocks


def _fetch_mermaid_png(mermaid_code: str) -> bytes | None:
    """Fetch a Mermaid diagram as PNG from mermaid.ink API."""
    import base64
    import urllib.request
    try:
        encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("ascii")
        url = f"https://mermaid.ink/img/{encoded}?type=png&bgColor=!white"
        req = urllib.request.Request(url, headers={"User-Agent": "LacunexAI/1.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            return resp.read()
    except Exception:
        return None


# Callout style configs for PDF
_PDF_CALLOUT_STYLES = {
    "key_point": {"icon": "[Key Point]", "bg": (239, 246, 255), "border": (59, 130, 246), "text": (30, 64, 175)},
    "important": {"icon": "[Important]", "bg": (255, 251, 235), "border": (245, 158, 11), "text": (146, 64, 14)},
    "warning":   {"icon": "[Warning]",   "bg": (254, 242, 242), "border": (239, 68, 68),  "text": (153, 27, 27)},
    "summary":   {"icon": "[Summary]",   "bg": (240, 253, 244), "border": (16, 185, 129), "text": (6, 95, 70)},
    "definition":{"icon": "[Definition]","bg": (245, 243, 255), "border": (139, 92, 246), "text": (91, 33, 182)},
    "example":   {"icon": "[Example]",   "bg": (236, 254, 255), "border": (6, 182, 212),  "text": (21, 94, 117)},
    "note":      {"icon": "[Note]",      "bg": (239, 246, 255), "border": (59, 130, 246), "text": (30, 64, 175)},
}


def generate_document_pdf(doc_json: dict, theme: str = "professional") -> bytes:
    """Generate a premium-quality themed PDF from structured document JSON."""
    from fpdf import FPDF

    title = doc_json.get("title", "Untitled Document")
    sections = doc_json.get("sections", [])
    toc = doc_json.get("table_of_contents", [])
    metadata = doc_json.get("metadata", {})

    # Theme accent colors
    ACCENTS = {
        "professional": (37, 99, 235),
        "dark": (139, 92, 246),
        "minimal": (75, 85, 99),
    }
    accent = ACCENTS.get(theme, ACCENTS["professional"])
    current_section_name = ""

    class DocPDF(FPDF):
        def header(self):
            if self.page_no() <= 1:
                return
            # Running header: section name (left), page (right)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(140, 140, 160)
            self.set_y(8)
            name = current_section_name[:60]
            self.cell(0, 4, _pdf_safe(name), align="L")
            self.cell(0, 4, f"Page {self.page_no()}", align="R")
            # Thin accent line under header
            self.set_draw_color(*accent)
            self.set_line_width(0.2)
            self.line(22, 14, 188, 14)
            self.set_y(18)

        def footer(self):
            self.set_y(-14)
            self.set_draw_color(200, 200, 215)
            self.set_line_width(0.15)
            self.line(22, self.get_y() - 2, 188, self.get_y() - 2)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(140, 140, 160)
            self.cell(0, 4, _pdf_safe(BRAND_FOOTER), align="C")

    pdf = DocPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)

    L, R = 22, 20
    PW = 210 - L - R  # 168mm usable

    pdf.set_left_margin(L)
    pdf.set_right_margin(R)
    pdf.set_top_margin(18)

    # ═══════════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ═══════════════════════════════════════════════════════════════════
    pdf.add_page()

    # Top accent bar (gradient feel with rect)
    pdf.set_fill_color(*accent)
    pdf.rect(0, 0, 210, 5, "F")

    # Side accent strip
    pdf.set_fill_color(*accent)
    pdf.rect(0, 0, 4, 297, "F")

    # Title area
    pdf.ln(60)
    pdf.set_font("Helvetica", "B", 26)
    pdf.set_text_color(15, 23, 42)
    safe_title = _pdf_safe(_clean(title))
    # Adjust font for very long titles
    if len(safe_title) > 60:
        pdf.set_font("Helvetica", "B", 22)
    pdf.multi_cell(PW, 11, safe_title, align="C")

    # Decorative separator
    pdf.ln(6)
    y = pdf.get_y()
    pdf.set_draw_color(*accent)
    pdf.set_line_width(0.8)
    mid = 105
    pdf.line(mid - 25, y, mid + 25, y)
    pdf.ln(8)

    # Subtitle
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(100, 116, 139)
    sec_count = metadata.get('total_sections', 0)
    page_est = metadata.get('total_pages_estimate', 1)
    subtitle = f"{sec_count} Sections  |  ~{page_est} Pages  |  {metadata.get('total_tables', 0)} Tables"
    pdf.cell(PW, 6, _pdf_safe(subtitle), align="C", ln=True)

    # Date
    pdf.ln(4)
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(148, 163, 184)
    pdf.cell(PW, 5, _pdf_safe(f"Generated on {_ts()}"), align="C", ln=True)

    # Brand
    pdf.ln(10)
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(*accent)
    pdf.cell(PW, 6, "LACUNEX AI", align="C", ln=True)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(148, 163, 184)
    pdf.cell(PW, 5, "Filling the gaps humans can't reach", align="C", ln=True)

    # Bottom accent strip on cover
    pdf.set_fill_color(*accent)
    pdf.rect(0, 292, 210, 5, "F")

    # ═══════════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════════
    if toc:
        pdf.add_page()
        current_section_name = "Table of Contents"

        pdf.set_font("Helvetica", "B", 18)
        pdf.set_text_color(15, 23, 42)
        pdf.cell(PW, 10, "Table of Contents", ln=True)

        # Accent underline
        pdf.set_draw_color(*accent)
        pdf.set_line_width(0.6)
        pdf.line(L, pdf.get_y() + 1, L + 50, pdf.get_y() + 1)
        pdf.ln(8)

        for idx, entry in enumerate(toc):
            # Main entry
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(*accent)
            pdf.cell(10, 7, f"{idx + 1}.")
            pdf.set_text_color(15, 23, 42)
            pdf.cell(PW - 10, 7, _pdf_safe(entry.get("title", "")[:70]), ln=True)

            # Sub-entries
            for child in entry.get("children", []):
                pdf.set_font("Helvetica", "", 9)
                pdf.set_text_color(100, 116, 139)
                pdf.cell(18, 5, "")
                pdf.cell(PW - 18, 5, _pdf_safe(child.get("title", "")[:65]), ln=True)

            pdf.ln(2)

        pdf.set_text_color(15, 23, 42)

    # ═══════════════════════════════════════════════════════════════════
    #  CONTENT SECTIONS
    # ═══════════════════════════════════════════════════════════════════
    for sec_idx, section in enumerate(sections):
        blocks = _doc_flatten_content(section)
        heading_text = section.get("heading", "Section")

        # Page break before each top-level section (except first)
        if sec_idx > 0 and section.get("level", 2) <= 2:
            pdf.add_page()

        current_section_name = heading_text

        for block in blocks:
            btype = block.get("type")

            # ── Blank line ────────────────────────────────────────────
            if btype == "blank":
                pdf.ln(2)

            # ── Heading ───────────────────────────────────────────────
            elif btype == "heading":
                lvl = block.get("level", 2)
                sizes = {1: 18, 2: 14, 3: 12, 4: 10}
                pdf.ln(5)

                if lvl <= 2:
                    # Accent divider before major headings
                    pdf.set_draw_color(*accent)
                    pdf.set_line_width(0.5)
                    pdf.line(L, pdf.get_y(), L + 40, pdf.get_y())
                    pdf.ln(3)

                pdf.set_font("Helvetica", "B", sizes.get(lvl, 10))
                pdf.set_text_color(15, 23, 42)
                pdf.multi_cell(PW, 7, _pdf_safe(_strip_md(block["text"])), align="L")

                if lvl <= 2:
                    pdf.set_draw_color(220, 220, 235)
                    pdf.set_line_width(0.15)
                    pdf.line(L, pdf.get_y() + 1, 210 - R, pdf.get_y() + 1)

                pdf.ln(3)

            # ── Bullet ────────────────────────────────────────────────
            elif btype == "bullet":
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 60)
                clean_text = _strip_md(block["text"])
                x_start = L + 6
                pdf.set_x(L + 2)
                pdf.set_text_color(*accent)
                pdf.cell(4, 5.5, _pdf_safe("\u2022"))
                pdf.set_text_color(40, 40, 60)
                pdf.set_x(x_start)
                old_lm = pdf.l_margin
                pdf.set_left_margin(x_start)
                pdf.multi_cell(PW - 6, 5.5, _pdf_safe(clean_text), align="L")
                pdf.set_left_margin(old_lm)

            # ── Numbered ──────────────────────────────────────────────
            elif btype == "numbered":
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 60)
                clean_text = _strip_md(block["text"])
                num_str = f"{block.get('number', '.')}."
                x_start = L + 8
                pdf.set_x(L + 2)
                pdf.set_text_color(*accent)
                pdf.cell(6, 5.5, _pdf_safe(num_str))
                pdf.set_text_color(40, 40, 60)
                pdf.set_x(x_start)
                old_lm = pdf.l_margin
                pdf.set_left_margin(x_start)
                pdf.multi_cell(PW - 8, 5.5, _pdf_safe(clean_text), align="L")
                pdf.set_left_margin(old_lm)

            # ── Text paragraph ────────────────────────────────────────
            elif btype == "text":
                clean_text = _strip_md(block["text"])
                if not clean_text.strip():
                    continue
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 60)
                pdf.multi_cell(PW, 5.8, _pdf_safe(clean_text), align="L")
                pdf.ln(1.5)

            # ── Code block ────────────────────────────────────────────
            elif btype == "code":
                lang = block.get("language", "")
                pdf.ln(3)
                # Code header bar
                pdf.set_fill_color(235, 238, 245)
                pdf.set_draw_color(200, 205, 220)
                y_top = pdf.get_y()
                pdf.rect(L, y_top, PW, 5, "F")
                if lang:
                    pdf.set_font("Helvetica", "I", 7)
                    pdf.set_text_color(120, 120, 150)
                    pdf.set_xy(L + PW - 25, y_top + 0.5)
                    pdf.cell(24, 4, _pdf_safe(lang.upper()), align="R")
                pdf.set_y(y_top + 5)

                # Code body
                pdf.set_fill_color(246, 248, 252)
                pdf.set_font("Courier", "", 8)
                pdf.set_text_color(55, 48, 120)
                for code_line in block.get("lines", []):
                    safe = _pdf_safe(code_line if code_line else " ")
                    pdf.set_x(L)
                    pdf.multi_cell(PW, 4.5, "  " + safe, fill=True, border=0, align="L")
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 60)
                pdf.ln(3)

            # ── Callout / Highlight Box ───────────────────────────────
            elif btype == "callout":
                ct = block.get("callout_type", "key_point")
                style = _PDF_CALLOUT_STYLES.get(ct, _PDF_CALLOUT_STYLES["key_point"])
                pdf.ln(3)

                y_start = pdf.get_y()
                # Background fill
                pdf.set_fill_color(*style["bg"])
                # Left accent border (simulated with rect)
                pdf.set_fill_color(*style["border"])
                pdf.rect(L, y_start, 1.5, 14, "F")
                pdf.set_fill_color(*style["bg"])
                pdf.rect(L + 1.5, y_start, PW - 1.5, 14, "F")

                # Icon/label
                pdf.set_font("Helvetica", "B", 8)
                pdf.set_text_color(*style["text"])
                pdf.set_xy(L + 5, y_start + 1.5)
                pdf.cell(30, 4, _pdf_safe(style["icon"]))

                # Content
                pdf.set_font("Helvetica", "", 9)
                pdf.set_xy(L + 5, y_start + 6)
                clean = _strip_md(block.get("text", ""))
                pdf.multi_cell(PW - 10, 4.5, _pdf_safe(clean[:300]), align="L")

                pdf.set_y(max(pdf.get_y(), y_start + 14) + 2)
                pdf.set_text_color(40, 40, 60)

            # ── Table ─────────────────────────────────────────────────
            elif btype == "table":
                table = block.get("data", {})
                headers = table.get("headers", [])
                rows = table.get("rows", [])
                if not headers:
                    continue

                col_count = len(headers)
                col_w = min(max(PW / col_count, 20), 60)
                total_w = col_w * col_count
                x_offset = L + max(0, (PW - total_w) / 2)  # center table
                pdf.ln(4)

                # Header row
                pdf.set_font("Helvetica", "B", 8)
                pdf.set_fill_color(30, 58, 95)
                pdf.set_text_color(255, 255, 255)
                pdf.set_x(x_offset)
                for h in headers:
                    pdf.cell(col_w, 7, _pdf_safe(h[:30]), border=0, fill=True, align="C")
                pdf.ln()

                # Data rows
                pdf.set_font("Helvetica", "", 8)
                for ri, row in enumerate(rows):
                    if ri % 2 == 0:
                        pdf.set_fill_color(248, 250, 252)
                    else:
                        pdf.set_fill_color(255, 255, 255)
                    pdf.set_text_color(40, 40, 60)
                    pdf.set_x(x_offset)
                    for ci, cell in enumerate(row):
                        pdf.cell(col_w, 6, _pdf_safe(str(cell)[:50]), border=0, fill=True, align="L")
                    pdf.ln()

                # Bottom line
                pdf.set_draw_color(200, 205, 220)
                pdf.set_line_width(0.15)
                pdf.line(x_offset, pdf.get_y(), x_offset + total_w, pdf.get_y())
                pdf.ln(4)

            # ── Diagram ───────────────────────────────────────────────
            elif btype == "diagram":
                diagram_title = block.get("title", "Diagram")
                diagram_code = block.get("code", "")
                pdf.ln(4)

                # Try to fetch rendered PNG
                png_data = _fetch_mermaid_png(diagram_code) if diagram_code else None

                if png_data and len(png_data) > 100:
                    # Embed real diagram image
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.set_text_color(*accent)
                    pdf.cell(PW, 5, _pdf_safe(f"Diagram: {diagram_title}"), align="C", ln=True)
                    pdf.ln(2)
                    try:
                        img_buf = io.BytesIO(png_data)
                        pdf.image(img_buf, x=L + 10, w=PW - 20)
                    except Exception:
                        pass  # Fallback: just show the title
                else:
                    # Fallback: show as labeled code block
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.set_text_color(*accent)
                    pdf.cell(PW, 5, _pdf_safe(f"Diagram: {diagram_title}"), align="C", ln=True)
                    pdf.set_font("Courier", "", 7)
                    pdf.set_text_color(80, 80, 100)
                    for dline in diagram_code.split("\n")[:20]:
                        pdf.cell(PW, 3.5, _pdf_safe("  " + dline), ln=True)

                pdf.set_text_color(40, 40, 60)
                pdf.ln(4)

    return bytes(pdf.output())



def generate_document_docx(doc_json: dict, theme: str = "professional") -> bytes:
    """Generate a themed DOCX from structured document JSON."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    title = doc_json.get("title", "Untitled Document")
    sections = doc_json.get("sections", [])
    toc = doc_json.get("table_of_contents", [])
    metadata = doc_json.get("metadata", {})

    def add_hr():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCDD")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def shade_para(p, color_hex: str):
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color_hex)
        pPr.append(shd)

    # ── Title Page ────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(120)
    tr = tp.add_run(_clean(title))
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(15, 23, 42)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(f"{metadata.get('total_sections', 0)} Sections | ~{metadata.get('total_pages_estimate', 1)} Pages")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(100, 116, 139)

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(f"Generated on {_ts()}")
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(148, 163, 184)
    dr.italic = True

    bp = doc.add_paragraph()
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run("LACUNEX AI")
    br.bold = True
    br.font.size = Pt(10)
    br.font.color.rgb = RGBColor(37, 99, 235)

    doc.add_page_break()

    # ── Table of Contents ─────────────────────────────────────────────────
    if toc:
        tcp = doc.add_paragraph()
        tcr = tcp.add_run("Table of Contents")
        tcr.bold = True
        tcr.font.size = Pt(18)
        tcr.font.color.rgb = RGBColor(15, 23, 42)
        tcp.paragraph_format.space_after = Pt(12)
        add_hr()

        for idx, entry in enumerate(toc):
            ep = doc.add_paragraph()
            nr = ep.add_run(f"{idx + 1}. ")
            nr.bold = True
            nr.font.size = Pt(11)
            nr.font.color.rgb = RGBColor(37, 99, 235)
            er = ep.add_run(entry.get("title", ""))
            er.font.size = Pt(11)
            er.font.color.rgb = RGBColor(15, 23, 42)
            ep.paragraph_format.space_after = Pt(4)

            for child in entry.get("children", []):
                cp = doc.add_paragraph()
                cp.paragraph_format.left_indent = Cm(1)
                cr = cp.add_run(f"   {child.get('title', '')}")
                cr.font.size = Pt(9)
                cr.font.color.rgb = RGBColor(100, 116, 139)
                cp.paragraph_format.space_after = Pt(2)

        doc.add_page_break()

    # ── Content ───────────────────────────────────────────────────────────
    for section in sections:
        blocks = _doc_flatten_content(section)

        for block in blocks:
            btype = block.get("type")

            if btype == "heading":
                lvl = block.get("level", 2)
                sizes = {1: 18, 2: 15, 3: 13, 4: 11}
                hp = doc.add_paragraph()
                hr_ = hp.add_run(block["text"])
                hr_.bold = True
                hr_.font.size = Pt(sizes.get(lvl, 11))
                hr_.font.color.rgb = RGBColor(15, 23, 42)
                hp.paragraph_format.space_before = Pt(12)
                hp.paragraph_format.space_after = Pt(6)
                if lvl <= 2:
                    add_hr()

            elif btype == "bullet":
                bp = doc.add_paragraph(style="List Bullet")
                for r in _inline_runs(block["text"]):
                    run = bp.add_run(r["text"])
                    run.bold = r["bold"]
                    run.italic = r["italic"]
                    run.font.size = Pt(10.5)
                bp.paragraph_format.space_after = Pt(2)

            elif btype == "numbered":
                np_ = doc.add_paragraph(style="List Number")
                for r in _inline_runs(block["text"]):
                    run = np_.add_run(r["text"])
                    run.bold = r["bold"]
                    run.italic = r["italic"]
                    run.font.size = Pt(10.5)
                np_.paragraph_format.space_after = Pt(2)

            elif btype == "text":
                p = doc.add_paragraph()
                for r in _inline_runs(block["text"]):
                    run = p.add_run(r["text"])
                    run.bold = r["bold"]
                    run.italic = r["italic"]
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(22, 22, 45)
                p.paragraph_format.space_after = Pt(4)

            elif btype == "code":
                for cl in block.get("lines", []):
                    cp = doc.add_paragraph()
                    cr = cp.add_run(cl if cl else " ")
                    cr.font.name = "Courier New"
                    cr.font.size = Pt(9)
                    cr.font.color.rgb = RGBColor(60, 50, 140)
                    cp.paragraph_format.left_indent = Cm(0.5)
                    cp.paragraph_format.space_after = Pt(0)
                    shade_para(cp, "F1F5F9")

            elif btype == "callout":
                cp = doc.add_paragraph()
                cr = cp.add_run(block.get("text", ""))
                cr.italic = True
                cr.font.size = Pt(10)
                cr.font.color.rgb = RGBColor(30, 64, 175)
                cp.paragraph_format.left_indent = Cm(0.5)
                shade_para(cp, "EFF6FF")

            elif btype == "table":
                table_data = block.get("data", {})
                headers = table_data.get("headers", [])
                rows = table_data.get("rows", [])
                if headers:
                    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
                    tbl.style = "Table Grid"
                    # Headers
                    for ci, h in enumerate(headers):
                        cell = tbl.rows[0].cells[ci]
                        cell.text = h
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.bold = True
                                r.font.size = Pt(9)
                                r.font.color.rgb = RGBColor(255, 255, 255)
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"), "1E40AF")
                        cell._tc.get_or_add_tcPr().append(shd)
                    # Rows
                    for ri, row in enumerate(rows):
                        for ci, val in enumerate(row):
                            if ci < len(tbl.columns):
                                cell = tbl.rows[ri + 1].cells[ci]
                                cell.text = str(val)
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.font.size = Pt(9)

            elif btype == "blank":
                doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────────────
    add_hr()
    fp = doc.add_paragraph()
    fr = fp.add_run(BRAND_FOOTER)
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(148, 163, 184)
    fr.italic = True
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_document_xlsx(doc_json: dict, theme: str = "professional") -> bytes:
    """Generate a themed XLSX from structured document JSON — extracts all tables."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    title = doc_json.get("title", "Untitled Document")
    sections = doc_json.get("sections", [])

    # Remove default sheet
    wb.remove(wb.active)

    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill("solid", fgColor="1E40AF")
    stripe_fill = PatternFill("solid", fgColor="F1F5F9")

    table_count = 0

    def _clean_xlsx_text(val):
        if val is None:
            return ""
        s = str(val)
        # Convert break tags to literal newlines, as wrap_text=True is enabled
        s = re.sub(r'(?i)<br\s*/?>', '\n', s)
        s = re.sub(r'(?i)<li>', '• ', s)
        s = re.sub(r'(?i)</li>', '\n', s)
        # Strip remaining HTML tags
        s = re.sub(r'<[^>]+>', '', s)
        # Remove control characters that crash openpyxl entirely
        s = re.sub(r'[\000-\010]|[\013-\014]|[\016-\037]', '', s)
        return s.strip()

    def process_section_tables(section, sheet_prefix=""):
        nonlocal table_count
        sec_name = section.get("heading", "Section")

        for table in section.get("tables", []):
            table_count += 1
            headers = table.get("headers", [])
            rows = table.get("rows", [])
            if not headers:
                continue

            # Create a sheet for each table
            safe_name = re.sub(r"[^\w\s]", "", sec_name)[:25].strip()
            sheet_name = f"{safe_name}_{table_count}" if safe_name else f"Table_{table_count}"
            ws = wb.create_sheet(title=sheet_name)

            # Title row
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
            ws.cell(row=1, column=1, value=sec_name)
            ws.cell(row=1, column=1).font = Font(bold=True, size=12, color="0F172A")
            ws.cell(row=1, column=1).alignment = Alignment(horizontal="left")

            # Headers
            for ci, h in enumerate(headers, 1):
                clean_h = _clean_xlsx_text(h)
                cell = ws.cell(row=3, column=ci, value=clean_h)
                cell.font = Font(bold=True, color="FFFFFF", size=10)
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = border

            # Data rows
            for ri, row in enumerate(rows, 4):
                for ci, val in enumerate(row, 1):
                    clean_val = _clean_xlsx_text(val)
                    cell = ws.cell(row=ri, column=ci, value=clean_val)
                    cell.border = border
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                    cell.font = Font(size=10)
                    if (ri - 4) % 2 == 1:
                        cell.fill = stripe_fill

            # Auto-width
            for ci in range(1, len(headers) + 1):
                col_header = _clean_xlsx_text(headers[ci - 1])
                col_vals = [_clean_xlsx_text(r[ci - 1]) if ci - 1 < len(r) else "" for r in rows]
                max_len = max([len(col_header)] + [len(v) for v in col_vals])
                ws.column_dimensions[ws.cell(row=3, column=ci).column_letter].width = min(max_len + 4, 50)

        # Recurse into subsections
        for sub in section.get("subsections", []):
            process_section_tables(sub)

    for section in sections:
        process_section_tables(section)

    # If no tables found, create a summary sheet
    if table_count == 0:
        ws = wb.create_sheet(title="Document Summary")
        ws.cell(row=1, column=1, value=title)
        ws.cell(row=1, column=1).font = Font(bold=True, size=14)
        ws.cell(row=3, column=1, value="Section")
        ws.cell(row=3, column=2, value="Content Preview")
        ws.cell(row=3, column=1).font = Font(bold=True, color="FFFFFF")
        ws.cell(row=3, column=2).font = Font(bold=True, color="FFFFFF")
        ws.cell(row=3, column=1).fill = header_fill
        ws.cell(row=3, column=2).fill = header_fill

        for idx, section in enumerate(sections, 4):
            ws.cell(row=idx, column=1, value=section.get("heading", ""))
            ws.cell(row=idx, column=1).font = Font(bold=True, size=10)
            content_preview = (section.get("content", "")[:200] + "...") if section.get("content") else ""
            ws.cell(row=idx, column=2, value=content_preview)
            ws.cell(row=idx, column=2).alignment = Alignment(wrap_text=True)

        ws.column_dimensions["A"].width = 30
        ws.column_dimensions["B"].width = 80

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════
#  v2 Premium Overrides — ReportLab PDF + Enhanced DOCX
# ═══════════════════════════════════════════════════════════════════════════
# These imports shadow the original functions above with the premium v2
# implementations (cover pages, diagrams, callout boxes, styled tables).
from services._export_v2 import (                   # noqa: E402
    generate_document_pdf,                           # noqa: F811
    generate_document_docx,                          # noqa: F811
    _doc_flatten_content,                            # noqa: F811
)

