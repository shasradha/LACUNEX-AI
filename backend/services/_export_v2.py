"""
LACUNEX AI — Premium Document Export Engine v2
Generates magazine-quality PDF (ReportLab Platypus) and DOCX (python-docx).
Shared design language: Navy #0a0a2e · Cyan #00d4ff · White · Light Gray.
"""

import io
import re

from services.export_service import (
    _pdf_safe, _clean, _strip_md, _ts,
    _inline_runs, _fetch_mermaid_png, BRAND_FOOTER,
)

# ═══════════════════════════════════════════════════════════════════════════
#  SHARED DESIGN LANGUAGE
# ═══════════════════════════════════════════════════════════════════════════

NAVY = "#0a0a2e"
CYAN = "#00d4ff"
WHITE = "#ffffff"
LIGHT = "#f0f4f8"
BODY_TEXT = "#2d2d2d"
DARK_NAVY = "#1a1a4e"
GRAY = "#555555"
ALT_ROW = "#f0f8ff"

_CALLOUT_STYLES = {
    "key_point":  {"label": "Key Point",  "bg": "#e0f7ff", "border": "#00d4ff", "tc": "#0066aa"},
    "warning":    {"label": "Warning",    "bg": "#fff8e1", "border": "#f59e0b", "tc": "#b45309"},
    "note":       {"label": "Note",       "bg": "#f5f5f5", "border": "#0a0a2e", "tc": "#333333"},
    "pro_tip":    {"label": "Pro Tip",    "bg": "#e8f5e9", "border": "#16a34a", "tc": "#166534"},
    "important":  {"label": "Important",  "bg": "#fff8e1", "border": "#f59e0b", "tc": "#b45309"},
    "summary":    {"label": "Summary",    "bg": "#e8f5e9", "border": "#16a34a", "tc": "#166534"},
    "definition": {"label": "Definition", "bg": "#f0f4ff", "border": "#6366f1", "tc": "#4338ca"},
    "example":    {"label": "Example",    "bg": "#ecfeff", "border": "#06b6d4", "tc": "#155e75"},
}


# ═══════════════════════════════════════════════════════════════════════════
#  CONTENT FLATTENER v2 — full callout / heading / table / code detection
# ═══════════════════════════════════════════════════════════════════════════

_CALLOUT_RE = re.compile(
    r'^>\s*\*\*(?:Key Point|Warning|Note|Pro Tip|Important|Tip|Example|'
    r'Summary|Definition|Caution)[:\s]*\*\*\s*(.+)$',
    re.I,
)
_CALLOUT_TYPE_RE = re.compile(
    r'^>\s*\*\*(Key Point|Warning|Note|Pro Tip|Important|Tip|Example|'
    r'Summary|Definition|Caution)',
    re.I,
)
_TYPE_MAP = {
    "key_point": "key_point", "warning": "warning", "note": "note",
    "pro_tip": "pro_tip", "important": "important", "tip": "key_point",
    "example": "example", "summary": "summary", "definition": "definition",
    "caution": "warning",
}


def _doc_flatten_content(section: dict, depth: int = 0) -> list[dict]:
    """Flatten a section tree into renderable blocks (v2 — enhanced detection)."""
    blocks = []
    level = section.get("level", 2)
    heading = section.get("heading", "")

    if heading:
        blocks.append({"type": "heading", "level": level, "text": heading})

    # ── Parse raw content text ────────────────────────────────────────────
    if section.get("content"):
        lines = section["content"].split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not stripped:
                blocks.append({"type": "blank"})
                i += 1
                continue

            # Callout patterns
            cm = _CALLOUT_RE.match(stripped)
            if cm:
                tm = _CALLOUT_TYPE_RE.match(stripped)
                ct = "key_point"
                if tm:
                    raw = tm.group(1).lower().replace(" ", "_")
                    ct = _TYPE_MAP.get(raw, "key_point")
                blocks.append({"type": "callout", "callout_type": ct,
                               "text": cm.group(1).strip()})
                i += 1
                continue

            # Heading in content
            hm = re.match(r"^(#{1,4})\s+(.*)", stripped)
            if hm:
                blocks.append({"type": "heading",
                               "level": len(hm.group(1)),
                               "text": hm.group(2).strip()})
                i += 1
                continue

            # Table detection
            if (stripped.startswith("|") and stripped.endswith("|")
                    and i + 1 < len(lines)
                    and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip())):
                headers = [c.strip() for c in stripped.strip("|").split("|")]
                i += 2
                rows = []
                while (i < len(lines)
                       and lines[i].strip().startswith("|")
                       and lines[i].strip().endswith("|")):
                    cells = [c.strip()
                             for c in lines[i].strip().strip("|").split("|")]
                    rows.append(cells)
                    i += 1
                blocks.append({"type": "table",
                               "data": {"headers": headers, "rows": rows}})
                continue

            # Code fence
            if stripped.startswith("```"):
                lang = stripped[3:].strip()
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    i += 1
                blocks.append({"type": "code", "lines": code_lines,
                               "language": lang})
                continue

            # Bullet
            bm = re.match(r"^[-*+]\s+(.*)", stripped)
            if bm:
                blocks.append({"type": "bullet",
                               "text": bm.group(1).strip()})
                i += 1
                continue

            # Numbered list
            nm = re.match(r"^(\d+)\.\s+(.*)", stripped)
            if nm:
                blocks.append({"type": "numbered",
                               "number": int(nm.group(1)),
                               "text": nm.group(2).strip()})
                i += 1
                continue

            # Plain text
            blocks.append({"type": "text", "text": stripped})
            i += 1

    # ── Structured elements from parser ───────────────────────────────────
    for table in section.get("tables", []):
        blocks.append({"type": "table", "data": table})
    for code in section.get("code_blocks", []):
        blocks.append({"type": "code",
                       "lines": code.get("code", "").split("\n"),
                       "language": code.get("language", "")})
    for hl in section.get("highlights", []):
        blocks.append({"type": "callout",
                       "callout_type": hl.get("type", "key_point"),
                       "text": hl.get("text", "")})
    for diag in section.get("diagrams", []):
        blocks.append({"type": "diagram",
                       "title": diag.get("title", "Diagram"),
                       "code": diag.get("code", "")})

    for sub in section.get("subsections", []):
        blocks.extend(_doc_flatten_content(sub, depth + 1))
    return blocks


# ═══════════════════════════════════════════════════════════════════════════
#  RICH-TEXT HELPER (ReportLab Paragraph XML)
# ═══════════════════════════════════════════════════════════════════════════

def _para_rich(text: str) -> str:
    """Convert markdown inline formatting → ReportLab Paragraph XML."""
    text = _pdf_safe(text)
    # Temporarily hide <br> tags from HTML escaping execution
    text = re.sub(r'(?i)<br\s*/?>', '{{BR_TAG_TEMP}}', text)
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    # Restore <br> tags as proper XML self-closing tags
    text = text.replace('{{BR_TAG_TEMP}}', '<br/>')
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    text = re.sub(r'`([^`]+)`',
                  r'<font face="Courier" size="9" color="#37306e">\1</font>',
                  text)
    return text


# ═══════════════════════════════════════════════════════════════════════════
#  DIAGRAM GENERATORS  (ReportLab Drawing / shapes)
# ═══════════════════════════════════════════════════════════════════════════

def _make_timeline(events, width=400):
    """Horizontal timeline with labeled milestone circles."""
    from reportlab.graphics.shapes import Drawing, Line, Circle, String
    from reportlab.lib.colors import HexColor
    d = Drawing(width, 80)
    if not events:
        return d
    n = min(len(events), 10)
    step = (width - 40) / max(n, 1)
    d.add(Line(20, 40, width - 20, 40,
               strokeColor=HexColor(CYAN), strokeWidth=2))
    for i, (label, desc) in enumerate(events[:n]):
        x = 20 + step * (i + 0.5)
        d.add(Circle(x, 40, 6, fillColor=HexColor(NAVY),
                     strokeColor=HexColor(CYAN), strokeWidth=2))
        d.add(String(x, 54, str(label)[:15], fontSize=7,
                     textAnchor='middle', fillColor=HexColor(CYAN)))
        d.add(String(x, 22, str(desc)[:20], fontSize=5.5,
                     textAnchor='middle', fillColor=HexColor(BODY_TEXT)))
    return d


def _make_bar_chart(data_pairs, chart_title="", width=400):
    """Vertical bar chart for comparisons."""
    from reportlab.graphics.shapes import Drawing, String
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.lib.colors import HexColor
    d = Drawing(width, 160)
    if not data_pairs:
        return d
    chart = VerticalBarChart()
    chart.x, chart.y = 50, 30
    chart.width, chart.height = width - 80, 100
    vals = [max(0, v) for _, v in data_pairs[:8]]
    chart.data = [vals or [1]]
    chart.categoryAxis.categoryNames = [str(l)[:12]
                                        for l, _ in data_pairs[:8]]
    chart.categoryAxis.labels.fontSize = 7
    chart.categoryAxis.labels.angle = 30
    chart.valueAxis.labels.fontSize = 7
    chart.bars[0].fillColor = HexColor(CYAN)
    chart.bars[0].strokeColor = HexColor(NAVY)
    d.add(chart)
    if chart_title:
        d.add(String(width / 2, 145, str(chart_title)[:50], fontSize=8,
                     textAnchor='middle', fillColor=HexColor(NAVY)))
    return d


def _make_architecture(layers, width=400):
    """Nested-rectangle layered architecture diagram."""
    from reportlab.graphics.shapes import Drawing, Rect, String
    from reportlab.lib.colors import HexColor
    n = min(len(layers), 6) or 1
    layers = layers or ["System"]
    h = n * 30 + 20
    d = Drawing(width, h)
    colors = [NAVY, DARK_NAVY, '#2a2a6e', '#3a3a8e', CYAN, '#66e5ff']
    for i, name in enumerate(layers[:n]):
        y = h - (i + 1) * 30
        margin = i * 15
        w = width - 2 * margin - 20
        d.add(Rect(10 + margin, y, w, 24,
                   fillColor=HexColor(colors[i % len(colors)]),
                   strokeColor=HexColor(WHITE), strokeWidth=1, rx=4))
        tc = WHITE if i < 3 else NAVY
        d.add(String(width / 2, y + 7, str(name)[:30], fontSize=8,
                     textAnchor='middle', fillColor=HexColor(tc)))
    return d


def _make_flowchart(steps, width=400):
    """Vertical flowchart with boxes and arrows."""
    from reportlab.graphics.shapes import Drawing, Rect, String, Line, Polygon
    from reportlab.lib.colors import HexColor
    n = min(len(steps), 8) or 1
    steps = steps or ["Start"]
    step_h, gap = 26, 15
    h = n * (step_h + gap) + 10
    d = Drawing(width, h)
    bw, cx = min(200, width - 60), width / 2
    for i, step in enumerate(steps[:n]):
        y = h - (i + 1) * (step_h + gap) + gap
        d.add(Rect(cx - bw / 2, y, bw, step_h,
                   fillColor=HexColor(LIGHT),
                   strokeColor=HexColor(CYAN), strokeWidth=1.5, rx=6))
        d.add(String(cx, y + 8, str(step)[:35], fontSize=8,
                     textAnchor='middle', fillColor=HexColor(NAVY)))
        if i < n - 1:
            d.add(Line(cx, y, cx, y - gap + 4,
                       strokeColor=HexColor(CYAN), strokeWidth=1.5))
            d.add(Polygon(
                [cx - 3, y - gap + 7, cx + 3, y - gap + 7, cx, y - gap + 2],
                fillColor=HexColor(CYAN), strokeColor=HexColor(CYAN)))
    return d


def _make_hierarchy(root, children, width=400):
    """Root → children tree layout."""
    from reportlab.graphics.shapes import Drawing, Rect, String, Line
    from reportlab.lib.colors import HexColor
    d = Drawing(width, 110)
    rw, rh = 130, 24
    rx, ry = width / 2 - rw / 2, 80
    d.add(Rect(rx, ry, rw, rh, fillColor=HexColor(NAVY),
               strokeColor=HexColor(CYAN), strokeWidth=1.5, rx=4))
    d.add(String(width / 2, ry + 7, str(root)[:20], fontSize=8,
                 textAnchor='middle', fillColor=HexColor(WHITE)))
    n = min(len(children), 5)
    if n:
        cw = min(80, (width - 40) / n - 10)
        total = n * cw + (n - 1) * 10
        sx, cy = (width - total) / 2, 15
        for i, ch in enumerate(children[:n]):
            cxi = sx + i * (cw + 10)
            d.add(Rect(cxi, cy, cw, rh, fillColor=HexColor(LIGHT),
                       strokeColor=HexColor(CYAN), strokeWidth=1, rx=4))
            d.add(String(cxi + cw / 2, cy + 7, str(ch)[:12], fontSize=6.5,
                         textAnchor='middle', fillColor=HexColor(NAVY)))
            d.add(Line(width / 2, ry, cxi + cw / 2, cy + rh,
                       strokeColor=HexColor('#cccccc'), strokeWidth=0.8))
    return d


# ── Auto-detection ────────────────────────────────────────────────────────

def _auto_detect_diagram_type(heading, content):
    c = f"{heading} {content}".lower()
    if any(k in c for k in ['timeline', 'history', 'evolution', 'year',
                             'era', 'century', 'generation', 'milestone']):
        return 'timeline'
    if any(k in c for k in ['compare', 'comparison', 'vs', 'versus',
                             'benchmark', 'performance', 'statistics',
                             'metric', 'score']):
        return 'bar_chart'
    if any(k in c for k in ['architecture', 'layer', 'stack', 'system',
                             'tier', 'infrastructure', 'component']):
        return 'architecture'
    if any(k in c for k in ['process', 'step', 'workflow', 'pipeline',
                             'flow', 'procedure', 'sequence', 'stage']):
        return 'flowchart'
    if any(k in c for k in ['hierarchy', 'tree', 'classification',
                             'taxonomy', 'category', 'type', 'kind']):
        return 'hierarchy'
    return None


def _generate_section_diagram(heading, content, width=400):
    """Auto-generate a relevant Drawing for a section."""
    dtype = _auto_detect_diagram_type(heading, content)
    if dtype is None:
        return None
    lines = [l.strip() for l in content.split('\n') if l.strip()]
    bullets = [re.sub(r'^[-*+]\s+', '', l)
               for l in lines if re.match(r'^[-*+]\s+', l)]
    numbered = [re.sub(r'^\d+\.\s+', '', l)
                for l in lines if re.match(r'^\d+\.\s+', l)]
    items = (bullets or numbered
             or [l for l in lines if len(l) > 5][:6]
             or [heading])
    if dtype == 'timeline':
        return _make_timeline(
            [(f"Phase {i+1}", it[:20]) for i, it in enumerate(items[:8])],
            width)
    if dtype == 'bar_chart':
        return _make_bar_chart(
            [(it[:15], 40 + i * 12) for i, it in enumerate(items[:8])],
            heading[:40], width)
    if dtype == 'architecture':
        return _make_architecture(items[:6], width)
    if dtype == 'flowchart':
        return _make_flowchart(items[:8], width)
    if dtype == 'hierarchy':
        return _make_hierarchy(heading[:20], items[:5], width)
    return None


# ═══════════════════════════════════════════════════════════════════════════
#  PDF GENERATOR — ReportLab Platypus
# ═══════════════════════════════════════════════════════════════════════════

def generate_document_pdf(doc_json: dict, theme: str = "professional") -> bytes:
    """Generate a premium PDF from structured document JSON."""
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, Flowable,
    )
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

    # Colours
    C_NAVY  = HexColor(NAVY)
    C_CYAN  = HexColor(CYAN)
    C_WHITE = HexColor(WHITE)
    C_BODY  = HexColor(BODY_TEXT)
    C_GRAY  = HexColor(GRAY)
    C_ALT   = HexColor(ALT_ROW)

    title_raw = doc_json.get("title", "Untitled Document")
    title = _pdf_safe(_clean(title_raw))
    sections = doc_json.get("sections", [])
    toc = doc_json.get("table_of_contents", [])
    meta = doc_json.get("metadata", {})

    page_w, page_h = A4
    LM, RM, TM, BM = 22*mm, 20*mm, 22*mm, 20*mm
    PW = page_w - LM - RM
    CH = page_h - TM - BM

    state = {"chapter": title}

    # ── Paragraph styles ──────────────────────────────────────────────────
    S = {
        'body': ParagraphStyle('body', fontName='Helvetica', fontSize=10.5,
                               textColor=C_BODY, leading=15,
                               alignment=TA_JUSTIFY, spaceAfter=3*mm),
        'h1': ParagraphStyle('h1', fontName='Helvetica-Bold', fontSize=20,
                             textColor=C_NAVY, spaceBefore=8*mm,
                             spaceAfter=4*mm, leading=24),
        'h2': ParagraphStyle('h2', fontName='Helvetica-Bold', fontSize=15,
                             textColor=HexColor(DARK_NAVY),
                             spaceBefore=6*mm, spaceAfter=3*mm, leading=19),
        'h3': ParagraphStyle('h3', fontName='Helvetica-Bold', fontSize=12,
                             textColor=C_CYAN, spaceBefore=4*mm,
                             spaceAfter=2*mm, leading=15),
        'h4': ParagraphStyle('h4', fontName='Helvetica-BoldOblique',
                             fontSize=10.5, textColor=C_GRAY,
                             spaceBefore=3*mm, spaceAfter=2*mm, leading=13),
        'bullet': ParagraphStyle('bullet', fontName='Helvetica',
                                 fontSize=10.5, textColor=C_BODY,
                                 leading=15, leftIndent=10*mm,
                                 firstLineIndent=-5*mm, spaceAfter=1.5*mm),
        'num': ParagraphStyle('num', fontName='Helvetica',
                              fontSize=10.5, textColor=C_BODY,
                              leading=15, leftIndent=10*mm,
                              firstLineIndent=-5*mm, spaceAfter=1.5*mm),
        'code': ParagraphStyle('code', fontName='Courier', fontSize=8.5,
                               textColor=HexColor('#37306e'), leading=11,
                               leftIndent=3*mm, spaceAfter=0),
    }

    # ── Custom flowables ──────────────────────────────────────────────────
    class ChapterDivider(Flowable):
        """Navy banner with chapter number + title."""
        def __init__(self, num, name, desc=""):
            Flowable.__init__(self)
            self.num = num
            self.name = _pdf_safe(name)[:55]
            self.desc = _pdf_safe(desc)[:90]
            self.width = PW
            self.height = 60*mm

        def draw(self):
            c = self.canv
            w, h = float(self.width), float(self.height)
            c.setFillColor(C_NAVY)
            c.rect(0, 0, w, h, fill=1, stroke=0)
            c.setFillColor(C_CYAN)
            c.setFont("Helvetica-Bold", 48)
            c.drawCentredString(w/2, h - 25*mm, str(self.num))
            c.setFillColor(C_WHITE)
            c.setFont("Helvetica-Bold", 20)
            c.drawCentredString(w/2, h - 45*mm, self.name)
            if self.desc:
                c.setFillColor(HexColor('#aaaaaa'))
                c.setFont("Helvetica", 9)
                c.drawCentredString(w/2, h - 58*mm, self.desc)

    class AccentBar(Flowable):
        """Cyan left-bar accent for major headings."""
        def __init__(self, w_mm, h_mm, color=C_CYAN):
            Flowable.__init__(self)
            self.width = w_mm*mm
            self.height = h_mm*mm
            self.color = color

        def draw(self):
            self.canv.setFillColor(self.color)
            self.canv.rect(0, 0, float(self.width), float(self.height),
                           fill=1, stroke=0)

    class ChapterMarker(Flowable):
        """Invisible — updates running-header chapter name."""
        def __init__(self, name):
            Flowable.__init__(self)
            self._name = name
            self.width = self.height = 0

        def draw(self):
            state["chapter"] = self._name

    # ── Page callbacks ────────────────────────────────────────────────────
    def _cover(canvas, doc):
        canvas.saveState()
        w, h = A4
        canvas.setFillColor(C_NAVY)
        canvas.rect(0, 0, w, h, fill=1, stroke=0)
        # cyan rules
        canvas.setStrokeColor(C_CYAN); canvas.setLineWidth(3)
        canvas.line(0, h - 5*mm, w, h - 5*mm)
        canvas.line(0, 5*mm, w, 5*mm)
        # branding
        canvas.setFillColor(C_CYAN)
        canvas.setFont("Helvetica-Bold", 14)
        canvas.drawCentredString(w/2, h - 35*mm, "L A C U N E X   A I")
        canvas.setFillColor(C_WHITE)
        canvas.setFont("Helvetica-Oblique", 10)
        canvas.drawCentredString(w/2, h - 46*mm,
                                 "Filling the gaps humans can't reach")
        # title
        fs = 28 if len(title) <= 40 else (22 if len(title) <= 60 else 18)
        canvas.setFont("Helvetica-Bold", fs)
        canvas.drawCentredString(w/2, h/2 + 15*mm, title[:80])
        canvas.setFillColor(C_CYAN)
        canvas.rect(w/2 - 30*mm, h/2 + 8*mm, 60*mm, 1.5*mm, fill=1, stroke=0)
        # metadata
        canvas.setFont("Helvetica", 10)
        sc = meta.get('total_sections', 0)
        pe = meta.get('total_pages_estimate', 1)
        tc = meta.get('total_tables', 0)
        canvas.drawCentredString(
            w/2, h/2 - 8*mm,
            f"{sc} Sections  |  ~{pe} Pages  |  {tc} Tables")
        canvas.setFillColor(HexColor('#888888'))
        canvas.setFont("Helvetica-Oblique", 9)
        canvas.drawCentredString(w/2, h/2 - 20*mm,
                                 _pdf_safe(f"Generated on {_ts()}"))
        canvas.setFillColor(HexColor('#777777'))
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(
            w/2, 18*mm,
            "(c) Generated by LACUNEX AI -- Shasradha Karmakar")
        canvas.restoreState()

    def _content(canvas, doc):
        canvas.saveState()
        w, h = A4
        # header
        canvas.setFont("Helvetica-Bold", 8); canvas.setFillColor(C_NAVY)
        
        short_title = title[:50] + ("..." if len(title) > 50 else "")
        header_text = f"LACUNEX AI | {short_title} | Page {doc.page}"

        canvas.drawCentredString(w/2, h - 14*mm, _pdf_safe(header_text))
        canvas.setStrokeColor(C_CYAN); canvas.setLineWidth(0.5)
        canvas.line(LM, h - 17*mm, w - RM, h - 17*mm)
        # footer
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(HexColor('#999999'))
        canvas.drawString(LM, 8*mm, _pdf_safe(title[:55]))
        canvas.drawRightString(w - RM, 8*mm, "(c) LACUNEX AI")
        canvas.setStrokeColor(HexColor('#cccccc')); canvas.setLineWidth(0.3)
        canvas.line(LM, 12*mm, w - RM, 12*mm)
        canvas.restoreState()

    # ── Build story ───────────────────────────────────────────────────────
    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=LM, rightMargin=RM,
                            topMargin=TM, bottomMargin=BM,
                            title=title)
    story = []

    # Cover page — design is drawn by onFirstPage callback;
    # use a tiny spacer (frame has ~12pt less than CH due to internal padding)
    story.append(Spacer(1, 1))
    story.append(PageBreak())

    # ── TOC (max 2 pages) ─────────────────────────────────────────────────
    if toc:
        story.append(ChapterMarker("Table of Contents"))
        toc_bar = Table(
            [[AccentBar(1.5, 6),
              Paragraph("Table of Contents", S['h1'])]],
            colWidths=[6*mm, PW - 6*mm])
        toc_bar.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(toc_bar)
        story.append(Spacer(1, 4*mm))

        total = sum(1 + len(e.get("children", [])) for e in toc)
        if total > 40:
            fsz, ssz, ld = 8, 7, 10
        elif total > 25:
            fsz, ssz, ld = 9, 8, 12
        else:
            fsz, ssz, ld = 11, 9, 16

        ch_s = ParagraphStyle('tch', fontName='Helvetica-Bold',
                              fontSize=fsz, textColor=C_NAVY,
                              leading=ld, spaceAfter=1*mm)
        sub_s = ParagraphStyle('tcs', fontName='Helvetica',
                               fontSize=ssz, textColor=C_GRAY,
                               leading=ld - 2, leftIndent=10*mm,
                               spaceAfter=0.5*mm)

        for idx, entry in enumerate(toc):
            t = _pdf_safe(entry.get("title", ""))[:70]
            dots = '.' * max(2, 70 - len(t) - len(str(idx+1)) - 4)
            story.append(Paragraph(
                f"<b>{idx+1}.</b>  {t}  "
                f"<font color='#cccccc'>{dots}</font>", ch_s))
            for child in entry.get("children", []):
                ct = _pdf_safe(child.get("title", ""))[:60]
                d2 = '.' * max(2, 60 - len(ct))
                story.append(Paragraph(
                    f"{ct}  <font color='#dddddd'>{d2}</font>", sub_s))
        story.append(PageBreak())

    # ── Content sections ──────────────────────────────────────────────────
    diagram_count = 0

    for sec_idx, section in enumerate(sections):
        heading_text = section.get("heading", "Section")
        desc = (section.get("content", "") or "").split("\n")[0][:80]

        if section.get("level", 2) <= 2:
            if sec_idx > 0:
                story.append(PageBreak())
            story.append(ChapterMarker(heading_text))
            story.append(ChapterDivider(sec_idx + 1, heading_text, desc))
            story.append(Spacer(1, 6*mm))

        blocks = _doc_flatten_content(section)

        for block in blocks:
            bt = block.get("type")

            if bt == "blank":
                story.append(Spacer(1, 2*mm))

            elif bt == "heading":
                lvl = min(block.get("level", 2), 4)
                txt = _para_rich(block.get("text", ""))
                if lvl <= 2:
                    story.append(AccentBar(1.5 if lvl == 1 else 1, 5))
                story.append(Paragraph(txt, S[f'h{lvl}']))

            elif bt == "bullet":
                txt = _para_rich(block.get("text", ""))
                story.append(Paragraph(
                    f"<font color='{CYAN}'>&#8226;</font>  {txt}",
                    S['bullet']))

            elif bt == "numbered":
                txt = _para_rich(block.get("text", ""))
                num = block.get("number", 1)
                story.append(Paragraph(
                    f"<font color='{CYAN}'><b>{num}.</b></font>  {txt}",
                    S['num']))

            elif bt == "text":
                txt = _para_rich(block.get("text", ""))
                if txt.strip():
                    story.append(Paragraph(txt, S['body']))

            elif bt == "code":
                lang = block.get("language", "")
                code_lines = block.get("lines", [])
                if lang:
                    story.append(Paragraph(
                        _pdf_safe(lang.upper()),
                        ParagraphStyle('lang', fontName='Helvetica-Oblique',
                                       fontSize=7, textColor=C_GRAY,
                                       spaceBefore=2*mm, spaceAfter=0)))
                safe_lines = []
                for cl in code_lines:
                    s = _pdf_safe(cl) if cl.strip() else " "
                    s = s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    safe_lines.append(s)
                code_text = "<br/>".join(safe_lines)
                ct = Table(
                    [[Paragraph(code_text, S['code'])]],
                    colWidths=[PW])
                ct.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), HexColor('#f6f8fc')),
                    ('BOX', (0, 0), (-1, -1), 0.5, HexColor('#e0e0e0')),
                    ('LEFTPADDING', (0, 0), (-1, -1), 3*mm),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 3*mm),
                    ('TOPPADDING', (0, 0), (-1, -1), 2*mm),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2*mm),
                ]))
                story.append(ct)
                story.append(Spacer(1, 3*mm))

            elif bt == "callout":
                ctype = block.get("callout_type", "key_point")
                cs = _CALLOUT_STYLES.get(ctype, _CALLOUT_STYLES["key_point"])
                txt = _para_rich(block.get("text", ""))[:500]
                lbl_s = ParagraphStyle('cl', fontName='Helvetica-Bold',
                                       fontSize=9,
                                       textColor=HexColor(cs["tc"]),
                                       spaceAfter=1*mm)
                bdy_s = ParagraphStyle('cb', fontName='Helvetica',
                                       fontSize=9, textColor=C_BODY,
                                       leading=13)
                inner = [[Paragraph(f"[{cs['label']}]", lbl_s)],
                         [Paragraph(txt, bdy_s)]]
                ct = Table(inner, colWidths=[PW - 4*mm])
                ct.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), HexColor(cs["bg"])),
                    ('LINEBEFORE', (0, 0), (0, -1), 4, HexColor(cs["border"])),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6*mm),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4*mm),
                    ('TOPPADDING', (0, 0), (0, 0), 3*mm),
                    ('BOTTOMPADDING', (-1, -1), (-1, -1), 3*mm),
                ]))
                story.append(Spacer(1, 2*mm))
                story.append(ct)
                story.append(Spacer(1, 2*mm))

            elif bt == "table":
                td = block.get("data", {})
                hdrs = td.get("headers", [])
                rows = td.get("rows", [])
                if not hdrs:
                    continue
                nc = len(hdrs)
                cw = PW / nc
                h_s = ParagraphStyle('th', fontName='Helvetica-Bold',
                                     fontSize=9, textColor=C_WHITE,
                                     alignment=TA_CENTER)
                c_s = ParagraphStyle('tc', fontName='Helvetica',
                                     fontSize=8.5, textColor=C_BODY,
                                     leading=12)
                data = [[Paragraph(_para_rich(str(h)), h_s)
                         for h in hdrs]]
                for row in rows:
                    data.append([Paragraph(_para_rich(str(v)), c_s)
                                 for v in row])
                tbl = Table(data, colWidths=[cw]*nc, repeatRows=1)
                ts_list = [
                    ('BACKGROUND', (0, 0), (-1, 0), C_NAVY),
                    ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#cccccc')),
                    ('TOPPADDING', (0, 0), (-1, -1), 2*mm),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2*mm),
                    ('LEFTPADDING', (0, 0), (-1, -1), 2*mm),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 2*mm),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]
                for ri in range(1, len(data)):
                    bg = C_ALT if ri % 2 == 1 else C_WHITE
                    ts_list.append(('BACKGROUND', (0, ri), (-1, ri), bg))
                tbl.setStyle(TableStyle(ts_list))
                story.append(Spacer(1, 3*mm))
                story.append(tbl)
                story.append(Spacer(1, 3*mm))

            elif bt == "diagram":
                dtitle = block.get("title", "Diagram")
                dcode = block.get("code", "")
                png = _fetch_mermaid_png(dcode) if dcode else None
                if png and len(png) > 100:
                    try:
                        from reportlab.platypus import Image as RLImage
                        story.append(Spacer(1, 3*mm))
                        story.append(Paragraph(
                            _pdf_safe(f"Diagram: {dtitle}"), S['h3']))
                        story.append(RLImage(io.BytesIO(png),
                                             width=PW - 20*mm,
                                             height=70*mm))
                        story.append(Spacer(1, 3*mm))
                    except Exception:
                        pass

        # Auto-generate diagram per section
        sec_content = section.get("content", "") or ""
        for sub in section.get("subsections", []):
            sec_content += " " + (sub.get("content", "") or "")
        dia = _generate_section_diagram(heading_text, sec_content,
                                        width=float(PW))
        if dia:
            diagram_count += 1
            fig_s = ParagraphStyle('fig_t', fontName='Helvetica-Bold',
                                   fontSize=9, textColor=C_CYAN,
                                   alignment=TA_CENTER, spaceAfter=2*mm)
            story.append(Spacer(1, 3*mm))
            story.append(Paragraph(
                _pdf_safe(f"Figure {diagram_count}: {heading_text[:45]}"),
                fig_s))
            story.append(dia)
            story.append(Spacer(1, 3*mm))

    # Ensure minimum 3 diagrams
    if diagram_count < 3 and sections:
        extras = [
            lambda: _make_flowchart(
                ["Start", "Analyze", "Process", "Execute", "Complete"],
                float(PW)),
            lambda: _make_bar_chart(
                [(s.get("heading", "")[:15],
                  max(10, len(s.get("content", ""))))
                 for s in sections[:8]] or [("A", 50)],
                "Section Content Overview", float(PW)),
            lambda: _make_architecture(
                [s.get("heading", "Layer")[:20]
                 for s in sections[:6]] or ["System"],
                float(PW)),
        ]
        fig_s = ParagraphStyle('fig2', fontName='Helvetica-Bold',
                               fontSize=9, textColor=C_CYAN,
                               alignment=TA_CENTER, spaceAfter=2*mm)
        for builder in extras:
            if diagram_count >= 3:
                break
            diagram_count += 1
            story.append(Spacer(1, 3*mm))
            story.append(Paragraph(f"Figure {diagram_count}", fig_s))
            story.append(builder())
            story.append(Spacer(1, 3*mm))

    pdf.build(story, onFirstPage=_cover, onLaterPages=_content)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════
#  DOCX GENERATOR — python-docx with proper Word styles
# ═══════════════════════════════════════════════════════════════════════════

def generate_document_docx(doc_json: dict, theme: str = "professional") -> bytes:
    """Generate a premium DOCX from structured document JSON."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    title_raw = doc_json.get("title", "Untitled Document")
    title = _clean(title_raw)
    sections_data = doc_json.get("sections", [])
    toc = doc_json.get("table_of_contents", [])
    meta = doc_json.get("metadata", {})
    CONTENT_W = 9360  # DXA = 6.5 inches

    # ── Page setup ────────────────────────────────────────────────────────
    for sect in doc.sections:
        sect.page_width = Inches(8.5)
        sect.page_height = Inches(11)
        sect.top_margin = Inches(1)
        sect.bottom_margin = Inches(1)
        sect.left_margin = Inches(1)
        sect.right_margin = Inches(1)
        sect.different_first_page_header_footer = True

    # ── Heading styles ────────────────────────────────────────────────────
    def _style_heading(sname, font_size, color_hex, bold=True, italic=False,
                       outline_level=None, border_color=None, border_sz=24):
        style = doc.styles[sname]
        style.font.name = "Arial"
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor.from_string(color_hex)
        pf = style.paragraph_format
        pf.space_before = Pt(int(font_size * 1.2))
        pf.space_after = Pt(int(font_size * 0.6))
        if outline_level is not None:
            pPr = style.element.get_or_add_pPr()
            for old in pPr.findall(qn('w:outlineLvl')):
                pPr.remove(old)
            ol = OxmlElement('w:outlineLvl')
            ol.set(qn('w:val'), str(outline_level))
            pPr.append(ol)
        if border_color:
            pPr = style.element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), str(border_sz))
            left.set(qn('w:color'), border_color)
            left.set(qn('w:space'), '8')
            pBdr.append(left)
            pPr.append(pBdr)

    _style_heading('Heading 1', 20, '0a0a2e', outline_level=0,
                   border_color='00d4ff', border_sz=24)
    _style_heading('Heading 2', 16, '1a1a4e', outline_level=1)
    _style_heading('Heading 3', 13, '0066cc', outline_level=2)
    _style_heading('Heading 4', 11, '555555', italic=True, outline_level=3)

    # Normal style
    ns = doc.styles['Normal']
    ns.font.name = "Arial"
    ns.font.size = Pt(11)
    ns.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    ns.paragraph_format.line_spacing = 1.5
    ns.paragraph_format.space_after = Pt(6)

    # ── Helpers ───────────────────────────────────────────────────────────
    def add_hr(color="CCCCDD"):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "1")
        b.set(qn("w:color"), color)
        pBdr.append(b)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

    def add_runs(para, text, base_size=11, base_color=None):
        if base_color is None:
            base_color = RGBColor(0x2d, 0x2d, 0x2d)
        for r in _inline_runs(text):
            run = para.add_run(r["text"])
            run.bold = r["bold"]
            run.italic = r["italic"]
            run.font.size = Pt(base_size)
            run.font.name = "Arial"
            if r["code"]:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x37, 0x30, 0x6e)
            else:
                run.font.color.rgb = base_color

    def add_page_number(paragraph):
        r1 = paragraph.add_run()
        fc1 = OxmlElement('w:fldChar')
        fc1.set(qn('w:fldCharType'), 'begin')
        r1._r.append(fc1)
        r2 = paragraph.add_run()
        it = OxmlElement('w:instrText')
        it.set(qn('xml:space'), 'preserve')
        it.text = ' PAGE '
        r2._r.append(it)
        r3 = paragraph.add_run()
        fc2 = OxmlElement('w:fldChar')
        fc2.set(qn('w:fldCharType'), 'end')
        r3._r.append(fc2)

    def shade_cell(cell, hex_no_hash):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_no_hash)
        tcPr.append(shd)

    def set_cell_borders_left(cell, color, sz=24):
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement("w:tcBorders")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), str(sz))
        left.set(qn("w:color"), color)
        left.set(qn("w:space"), "0")
        tcB.append(left)
        for side in ["top", "right", "bottom"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:color"), "auto")
            tcB.append(el)
        tcPr.append(tcB)

    def set_cell_margins(cell, top=120, bot=120, left=180, right=180):
        tcPr = cell._tc.get_or_add_tcPr()
        m = OxmlElement("w:tcMar")
        for side, val in [("top", top), ("bottom", bot),
                          ("left", left), ("right", right)]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")
            m.append(el)
        tcPr.append(m)

    def add_callout(ctype, text):
        cs = _CALLOUT_STYLES.get(ctype, _CALLOUT_STYLES["key_point"])
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        p1 = cell.paragraphs[0]
        r1 = p1.add_run(f"[{cs['label']}]")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = "Arial"
        r1.font.color.rgb = RGBColor.from_string(cs["tc"].lstrip('#'))
        p2 = cell.add_paragraph()
        add_runs(p2, text, base_size=10)
        shade_cell(cell, cs["bg"].lstrip('#'))
        set_cell_borders_left(cell, cs["border"].lstrip('#'), 24)
        set_cell_margins(cell)

    def add_styled_table(headers, rows):
        nc = len(headers)
        if nc == 0:
            return
        tbl = doc.add_table(rows=1 + len(rows), cols=nc)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = 'Table Grid'
        for ci, h in enumerate(headers):
            cell = tbl.rows[0].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(h))
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = "Arial"
            r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            shade_cell(cell, NAVY.lstrip('#'))
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                if ci >= nc:
                    break
                cell = tbl.rows[ri + 1].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(str(val))
                r.font.size = Pt(9.5)
                r.font.name = "Arial"
                if ri % 2 == 0:
                    shade_cell(cell, ALT_ROW.lstrip('#'))

    # ── Header / Footer ──────────────────────────────────────────────────
    sect = doc.sections[0]

    # Default header (pages 2+)
    hdr = sect.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    # Tab stops via XML for portability
    hpPr = hp._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab_c = OxmlElement('w:tab')
    tab_c.set(qn('w:val'), 'center')
    tab_c.set(qn('w:pos'), '4680')
    tabs.append(tab_c)
    tab_r = OxmlElement('w:tab')
    tab_r.set(qn('w:val'), 'right')
    tab_r.set(qn('w:pos'), '9360')
    tabs.append(tab_r)
    hpPr.append(tabs)
    r1 = hp.add_run("LACUNEX AI")
    r1.bold = True
    r1.font.size = Pt(8)
    r1.font.name = "Arial"
    r1.font.color.rgb = RGBColor(0x0a, 0x0a, 0x2e)
    hp.add_run("\t")
    r2 = hp.add_run(title[:40])
    r2.font.size = Pt(8)
    r2.font.name = "Arial"
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    hp.add_run("\t")
    add_page_number(hp)
    # Cyan bottom border on header
    hBdr = OxmlElement("w:pBdr")
    hBot = OxmlElement("w:bottom")
    hBot.set(qn("w:val"), "single")
    hBot.set(qn("w:sz"), "6")
    hBot.set(qn("w:color"), "00d4ff")
    hBot.set(qn("w:space"), "1")
    hBdr.append(hBot)
    hpPr.append(hBdr)

    # Default footer (pages 2+)
    ftr = sect.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fpPr = fp._p.get_or_add_pPr()
    ftabs = OxmlElement('w:tabs')
    ftab_r = OxmlElement('w:tab')
    ftab_r.set(qn('w:val'), 'right')
    ftab_r.set(qn('w:pos'), '9360')
    ftabs.append(ftab_r)
    fpPr.append(ftabs)
    fr1 = fp.add_run(title[:45])
    fr1.font.size = Pt(7)
    fr1.font.name = "Arial"
    fr1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fp.add_run("\t")
    fr2 = fp.add_run("(c) LACUNEX AI")
    fr2.font.size = Pt(7)
    fr2.font.name = "Arial"
    fr2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fBdr = OxmlElement("w:pBdr")
    fTop = OxmlElement("w:top")
    fTop.set(qn("w:val"), "single")
    fTop.set(qn("w:sz"), "4")
    fTop.set(qn("w:color"), "cccccc")
    fTop.set(qn("w:space"), "1")
    fBdr.append(fTop)
    fpPr.append(fBdr)

    # ── Cover Page ────────────────────────────────────────────────────────
    bp = doc.add_paragraph()
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(24)
    br_ = bp.add_run("L A C U N E X   A I")
    br_.bold = True
    br_.font.size = Pt(14)
    br_.font.name = "Arial"
    br_.font.color.rgb = RGBColor(0x00, 0xd4, 0xff)

    tp_tag = doc.add_paragraph()
    tp_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr_tag = tp_tag.add_run("Filling the gaps humans can't reach")
    tr_tag.italic = True
    tr_tag.font.size = Pt(10)
    tr_tag.font.name = "Arial"
    tr_tag.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph().paragraph_format.space_before = Pt(60)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.name = "Arial"
    tr.font.color.rgb = RGBColor(0x0a, 0x0a, 0x2e)
    tp.paragraph_format.space_after = Pt(8)

    # Cyan bar (1-row table)
    bar_t = doc.add_table(rows=1, cols=1)
    bar_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    bar_c = bar_t.rows[0].cells[0]
    bar_c.text = ""
    shade_cell(bar_c, "00d4ff")
    trEl = bar_t.rows[0]._tr
    trPr = trEl.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), '60')
    trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), '4000')
    tcW.set(qn('w:type'), 'dxa')
    bar_c._tc.get_or_add_tcPr().append(tcW)

    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(16)
    sc = meta.get('total_sections', 0)
    pe = meta.get('total_pages_estimate', 1)
    tc = meta.get('total_tables', 0)
    mr_ = mp.add_run(f"{sc} Sections  |  ~{pe} Pages  |  {tc} Tables")
    mr_.font.size = Pt(10)
    mr_.font.name = "Arial"
    mr_.font.color.rgb = RGBColor(0x00, 0xd4, 0xff)

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr_ = dp.add_run(f"Generated on {_ts()}")
    dr_.italic = True
    dr_.font.size = Pt(9)
    dr_.font.name = "Arial"
    dr_.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(40)
    cr_ = cp.add_run("(c) Generated by LACUNEX AI -- Shasradha Karmakar")
    cr_.font.size = Pt(8)
    cr_.font.name = "Arial"
    cr_.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ── Table of Contents ─────────────────────────────────────────────────
    if toc:
        doc.add_heading("Table of Contents", level=1)
        add_hr("00d4ff")

        total = sum(1 + len(e.get("children", [])) for e in toc)
        ch_sz = 11 if total <= 25 else (9 if total <= 40 else 8)
        sub_sz = max(7, ch_sz - 1)

        for idx, entry in enumerate(toc):
            ep = doc.add_paragraph()
            ep.paragraph_format.space_after = Pt(2)
            nr = ep.add_run(f"{idx + 1}.  ")
            nr.bold = True
            nr.font.size = Pt(ch_sz)
            nr.font.name = "Arial"
            nr.font.color.rgb = RGBColor(0x00, 0xd4, 0xff)
            et = entry.get("title", "")[:70]
            dots = "." * max(2, 65 - len(et))
            er = ep.add_run(et)
            er.font.size = Pt(ch_sz)
            er.font.name = "Arial"
            er.font.color.rgb = RGBColor(0x0a, 0x0a, 0x2e)
            dr2 = ep.add_run(f"  {dots}")
            dr2.font.size = Pt(ch_sz)
            dr2.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

            for child in entry.get("children", []):
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Cm(1)
                sp.paragraph_format.space_after = Pt(1)
                ct = child.get("title", "")[:60]
                d2 = "." * max(2, 55 - len(ct))
                sr = sp.add_run(ct)
                sr.font.size = Pt(sub_sz)
                sr.font.name = "Arial"
                sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                sd = sp.add_run(f"  {d2}")
                sd.font.size = Pt(sub_sz)
                sd.font.color.rgb = RGBColor(0xdd, 0xdd, 0xdd)

        doc.add_page_break()

    # ── Content ───────────────────────────────────────────────────────────
    for sec_idx, section in enumerate(sections_data):
        heading_text = section.get("heading", "Section")

        # Chapter divider banner
        if section.get("level", 2) <= 2:
            if sec_idx > 0:
                doc.add_page_break()
            div_t = doc.add_table(rows=1, cols=1)
            div_t.alignment = WD_TABLE_ALIGNMENT.CENTER
            div_c = div_t.rows[0].cells[0]
            div_c.text = ""
            shade_cell(div_c, NAVY.lstrip('#'))
            set_cell_margins(div_c, 300, 300, 200, 200)
            np_ = div_c.paragraphs[0]
            np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
            n_r = np_.add_run(str(sec_idx + 1))
            n_r.bold = True
            n_r.font.size = Pt(36)
            n_r.font.name = "Arial"
            n_r.font.color.rgb = RGBColor(0x00, 0xd4, 0xff)
            tp2 = div_c.add_paragraph()
            tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t_r = tp2.add_run(heading_text[:55])
            t_r.bold = True
            t_r.font.size = Pt(20)
            t_r.font.name = "Arial"
            t_r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            doc.add_paragraph()

        blocks = _doc_flatten_content(section)

        for block in blocks:
            bt = block.get("type")

            if bt == "blank":
                doc.add_paragraph().paragraph_format.space_after = Pt(1)

            elif bt == "heading":
                lvl = min(block.get("level", 2), 4)
                doc.add_heading(block.get("text", ""), level=lvl)

            elif bt == "bullet":
                bp2 = doc.add_paragraph(style="List Bullet")
                add_runs(bp2, block.get("text", ""))
                bp2.paragraph_format.space_after = Pt(2)

            elif bt == "numbered":
                np2 = doc.add_paragraph(style="List Number")
                add_runs(np2, block.get("text", ""))
                np2.paragraph_format.space_after = Pt(2)

            elif bt == "text":
                txt = block.get("text", "")
                if txt.strip():
                    p = doc.add_paragraph()
                    add_runs(p, txt)
                    p.paragraph_format.space_after = Pt(4)

            elif bt == "code":
                lang = block.get("language", "")
                if lang:
                    lp = doc.add_paragraph()
                    lr = lp.add_run(lang.upper())
                    lr.italic = True
                    lr.font.size = Pt(7)
                    lr.font.name = "Arial"
                    lr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                    lp.paragraph_format.space_after = Pt(0)
                for cl in block.get("lines", []):
                    cp2 = doc.add_paragraph()
                    cr = cp2.add_run(cl if cl else " ")
                    cr.font.name = "Courier New"
                    cr.font.size = Pt(9)
                    cr.font.color.rgb = RGBColor(0x37, 0x30, 0x6e)
                    cp2.paragraph_format.left_indent = Cm(0.5)
                    cp2.paragraph_format.space_after = Pt(0)
                    cpPr = cp2._p.get_or_add_pPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "F1F5F9")
                    cpPr.append(shd)
                doc.add_paragraph().paragraph_format.space_after = Pt(4)

            elif bt == "callout":
                add_callout(block.get("callout_type", "key_point"),
                            block.get("text", ""))

            elif bt == "table":
                td = block.get("data", {})
                hdrs = td.get("headers", [])
                rows = td.get("rows", [])
                if hdrs:
                    add_styled_table(hdrs, rows)

    # ── Document footer ───────────────────────────────────────────────────
    add_hr()
    fp_end = doc.add_paragraph()
    fr_end = fp_end.add_run(BRAND_FOOTER)
    fr_end.font.size = Pt(8)
    fr_end.font.name = "Arial"
    fr_end.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    fr_end.italic = True
    fp_end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
