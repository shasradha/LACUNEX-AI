"""
Test script for LACUNEX AI v2 document export.
Generates a 5-section test document and exports to PDF + DOCX.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from services.export_service import generate_document_pdf, generate_document_docx

SAMPLE_DOC = {
    "type": "document",
    "title": "Artificial Intelligence in Modern Healthcare",
    "generated_at": "2026-04-09T12:00:00Z",
    "table_of_contents": [
        {"title": "Introduction to AI in Healthcare", "level": 1, "index": 0,
         "children": [
             {"title": "Historical Timeline", "level": 3, "parent_index": 0, "index": 0},
             {"title": "Key Milestones", "level": 3, "parent_index": 0, "index": 1},
         ]},
        {"title": "Machine Learning Pipeline for Diagnostics", "level": 2, "index": 1,
         "children": [
             {"title": "Data Collection Steps", "level": 3, "parent_index": 1, "index": 0},
             {"title": "Model Architecture Layers", "level": 3, "parent_index": 1, "index": 1},
         ]},
        {"title": "Performance Comparison of AI Models", "level": 2, "index": 2,
         "children": [
             {"title": "Benchmark Results", "level": 3, "parent_index": 2, "index": 0},
         ]},
        {"title": "System Architecture Overview", "level": 2, "index": 3,
         "children": [
             {"title": "Infrastructure Components", "level": 3, "parent_index": 3, "index": 0},
         ]},
        {"title": "Classification Hierarchy of AI Techniques", "level": 2, "index": 4,
         "children": [
             {"title": "Supervised vs Unsupervised", "level": 3, "parent_index": 4, "index": 0},
         ]},
    ],
    "sections": [
        {
            "heading": "Introduction to AI in Healthcare",
            "level": 1,
            "type": "introduction",
            "content": (
                "Artificial Intelligence has revolutionized healthcare over the past decade. "
                "From diagnostic imaging to drug discovery, AI systems are transforming how "
                "medical professionals approach patient care.\n\n"
                "The evolution and timeline of AI in healthcare spans several key milestones:\n"
                "- 2010: Early neural network experiments in radiology\n"
                "- 2015: Deep learning breakthroughs in pathology\n"
                "- 2018: FDA approval of first AI diagnostic tool\n"
                "- 2020: AI-powered COVID-19 screening systems\n"
                "- 2023: Generative AI for personalized treatment plans\n"
                "- 2025: Real-time AI surgical assistants\n\n"
                "> **Key Point:** AI is not replacing doctors but augmenting their capabilities "
                "with data-driven insights that improve accuracy and speed.\n\n"
                "### Historical Timeline\n\n"
                "The history of AI in medicine dates back to the 1970s with expert systems like MYCIN.\n\n"
                "### Key Milestones\n\n"
                "Several breakthrough moments defined the trajectory of medical AI."
            ),
            "subsections": [],
            "tables": [],
            "code_blocks": [],
            "highlights": [
                {"type": "key_point", "text": "AI reduces diagnostic errors by up to 30% in clinical trials."},
            ],
            "diagrams": [],
            "summary": "",
            "practice_questions": [],
            "diagram_candidate": True,
        },
        {
            "heading": "Machine Learning Pipeline for Diagnostics",
            "level": 2,
            "type": "content",
            "content": (
                "Building a diagnostic ML pipeline involves several sequential steps "
                "that must be executed in the correct order.\n\n"
                "The process flow for building a diagnostic pipeline:\n"
                "- Collect patient data from EHR systems\n"
                "- Clean and preprocess medical records\n"
                "- Engineer features from clinical variables\n"
                "- Train classification models\n"
                "- Validate against held-out test sets\n"
                "- Deploy to clinical workflow\n\n"
                "> **Warning:** Always validate AI models against diverse patient populations "
                "to avoid bias in diagnostic predictions.\n\n"
                "### Data Collection Steps\n\n"
                "Data must be sourced from multiple hospitals to ensure diversity.\n\n"
                "### Model Architecture Layers\n\n"
                "Modern diagnostic models use multi-layer architectures:\n"
                "- Input layer for raw patient data\n"
                "- Convolutional layers for image features\n"
                "- Attention layers for clinical notes\n"
                "- Dense layers for classification\n"
                "- Output layer with probability scores"
            ),
            "subsections": [],
            "tables": [],
            "code_blocks": [
                {
                    "language": "python",
                    "code": "import tensorflow as tf\n\nmodel = tf.keras.Sequential([\n    tf.keras.layers.Conv2D(64, 3, activation='relu'),\n    tf.keras.layers.MaxPooling2D(),\n    tf.keras.layers.Flatten(),\n    tf.keras.layers.Dense(128, activation='relu'),\n    tf.keras.layers.Dense(10, activation='softmax'),\n])",
                    "line_count": 8,
                }
            ],
            "highlights": [],
            "diagrams": [],
            "summary": "",
            "practice_questions": [],
            "diagram_candidate": True,
        },
        {
            "heading": "Performance Comparison of AI Models",
            "level": 2,
            "type": "content",
            "content": (
                "Comparing different AI model architectures reveals significant performance "
                "differences across various diagnostic tasks.\n\n"
                "The benchmark comparison shows CNN models outperforming traditional methods "
                "in image-based diagnostics, while transformer models excel in clinical NLP.\n\n"
                "> **Note:** Performance metrics should always be evaluated in the context "
                "of the specific clinical use case.\n\n"
                "### Benchmark Results\n\n"
                "The following table summarizes benchmark results across different models:"
            ),
            "subsections": [],
            "tables": [
                {
                    "headers": ["Model", "Accuracy", "Precision", "Recall", "F1 Score"],
                    "rows": [
                        ["CNN-ResNet", "94.2%", "93.8%", "94.5%", "94.1%"],
                        ["Vision Transformer", "95.1%", "94.9%", "95.3%", "95.1%"],
                        ["Random Forest", "87.3%", "86.9%", "87.7%", "87.3%"],
                        ["Logistic Regression", "82.1%", "81.5%", "82.8%", "82.1%"],
                        ["XGBoost", "91.7%", "91.2%", "92.1%", "91.7%"],
                    ],
                    "row_count": 5,
                    "col_count": 5,
                }
            ],
            "code_blocks": [],
            "highlights": [
                {"type": "important", "text": "Vision Transformers achieve the highest accuracy at 95.1%."},
            ],
            "diagrams": [],
            "summary": "",
            "practice_questions": [],
            "diagram_candidate": True,
        },
        {
            "heading": "System Architecture Overview",
            "level": 2,
            "type": "content",
            "content": (
                "The healthcare AI system architecture consists of multiple layers and "
                "infrastructure components working together.\n\n"
                "The system uses a multi-tier architecture:\n"
                "- User Interface Layer\n"
                "- API Gateway Layer\n"
                "- Application Logic Layer\n"
                "- Machine Learning Engine Layer\n"
                "- Data Storage Layer\n"
                "- Infrastructure Layer\n\n"
                "> **Pro Tip:** Use containerized microservices for each layer to enable "
                "independent scaling and fault isolation.\n\n"
                "### Infrastructure Components\n\n"
                "Each component runs in its own Docker container orchestrated by Kubernetes."
            ),
            "subsections": [],
            "tables": [
                {
                    "headers": ["Component", "Technology", "Purpose"],
                    "rows": [
                        ["API Gateway", "Kong / Nginx", "Request routing and rate limiting"],
                        ["ML Engine", "TensorFlow Serving", "Model inference at scale"],
                        ["Database", "PostgreSQL + Redis", "Persistent and cache storage"],
                        ["Message Queue", "RabbitMQ", "Async task processing"],
                    ],
                    "row_count": 4,
                    "col_count": 3,
                }
            ],
            "code_blocks": [],
            "highlights": [],
            "diagrams": [],
            "summary": "",
            "practice_questions": [],
            "diagram_candidate": True,
        },
        {
            "heading": "Classification Hierarchy of AI Techniques",
            "level": 2,
            "type": "content",
            "content": (
                "AI techniques in healthcare follow a clear taxonomy and classification "
                "hierarchy from broad categories to specific implementations.\n\n"
                "The hierarchy tree of AI technique types:\n"
                "- Supervised Learning\n"
                "- Unsupervised Learning\n"
                "- Reinforcement Learning\n"
                "- Semi-supervised Learning\n"
                "- Self-supervised Learning\n\n"
                "### Supervised vs Unsupervised\n\n"
                "Supervised learning requires labeled data while unsupervised learning "
                "discovers patterns in unlabeled datasets.\n\n"
                "> **Key Point:** Choose the learning paradigm based on data availability "
                "and the specific clinical task requirements."
            ),
            "subsections": [],
            "tables": [],
            "code_blocks": [],
            "highlights": [],
            "diagrams": [],
            "summary": "",
            "practice_questions": [],
            "diagram_candidate": True,
        },
    ],
    "metadata": {
        "total_sections": 5,
        "total_subsections": 6,
        "total_tables": 2,
        "total_code_blocks": 1,
        "total_highlights": 2,
        "total_pages_estimate": 12,
        "total_diagram_candidates": 5,
        "theme": "professional",
        "character_count": 5000,
        "word_count": 800,
    },
}


def main():
    print("=" * 60)
    print("  LACUNEX AI — Export v2 Test Suite")
    print("=" * 60)

    out_dir = os.path.join(os.path.dirname(__file__), "test_output")
    os.makedirs(out_dir, exist_ok=True)

    # ── PDF ────────────────────────────────────────────────────────────────
    print("\n[1/2] Generating PDF...")
    try:
        pdf_bytes = generate_document_pdf(SAMPLE_DOC)
        pdf_path = os.path.join(out_dir, "test_export.pdf")
        with open(pdf_path, "wb") as f:
            f.write(pdf_bytes)
        print(f"  OK  PDF saved  ({len(pdf_bytes):,} bytes) -> {pdf_path}")

        # Page count — count /Type /Page (not /Pages) in raw PDF bytes
        import re as _re
        with open(pdf_path, "rb") as pf:
            raw = pf.read()
        page_count = len(_re.findall(rb'/Type\s*/Page[^s]', raw))
        print(f"  OK  PDF pages:  {page_count}")

        assert page_count >= 3, "PDF should have at least 3 pages (cover + TOC + content)"
        print(f"  OK  PDF has cover + TOC + content pages")

    except Exception as e:
        print(f"  FAIL  PDF generation error: {e}")
        import traceback; traceback.print_exc()
        return 1

    # ── DOCX ──────────────────────────────────────────────────────────────
    print("\n[2/2] Generating DOCX...")
    try:
        docx_bytes = generate_document_docx(SAMPLE_DOC)
        docx_path = os.path.join(out_dir, "test_export.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        print(f"  OK  DOCX saved ({len(docx_bytes):,} bytes) -> {docx_path}")

        # Validate DOCX
        from docx import Document
        doc = Document(docx_path)
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        print(f"  OK  DOCX paragraphs: {para_count}")
        print(f"  OK  DOCX tables:     {table_count}")

        # Check headings are styled (not raw markdown)
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        print(f"  OK  DOCX styled headings: {len(headings)}")
        assert len(headings) > 0, "DOCX should have styled headings"

    except Exception as e:
        print(f"  FAIL  DOCX generation error: {e}")
        import traceback; traceback.print_exc()
        return 1

    # ── Summary ───────────────────────────────────────────────────────────
    print("\n" + "=" * 60)
    print("  ALL TESTS PASSED")
    print(f"  PDF:  {page_count} pages, {len(pdf_bytes):,} bytes")
    print(f"  DOCX: {para_count} paragraphs, {table_count} tables, {len(headings)} headings")
    print(f"  Files in: {out_dir}")
    print("=" * 60)
    return 0


if __name__ == "__main__":
    sys.exit(main())
